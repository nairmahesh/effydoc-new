import requests
import json
import uuid
import os
from datetime import datetime
import time

# Get the backend URL from the frontend .env file
BACKEND_URL = "https://7056c40d-c306-42f0-bfc6-379013315fc8.preview.emergentagent.com/api"

def test_document_viewer_functionality():
    """Test the document viewer functionality with focus on formatting preservation"""
    print("Starting document viewer functionality tests...")
    
    # Test variables
    base_url = BACKEND_URL
    headers = {"Content-Type": "application/json"}
    test_user_email = f"test.user.{uuid.uuid4()}@example.com"
    test_user_password = "SecurePassword123!"
    
    # 1. Register a test user
    print("\n1. Registering test user...")
    user_data = {
        "email": test_user_email,
        "full_name": "Test User",
        "role": "editor",
        "organization": "Test Organization",
        "password": test_user_password
    }
    
    response = requests.post(f"{base_url}/auth/register", json=user_data)
    assert response.status_code == 200, f"Failed to register test user: {response.text}"
    data = response.json()
    
    # Save token for subsequent tests
    access_token = data["access_token"]
    headers["Authorization"] = f"Bearer {access_token}"
    print("✅ User registration successful")
    
    # 2. Create a complex DOCX file with rich formatting
    print("\n2. Creating complex DOCX file with rich formatting...")
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        docx_path = "/tmp/complex_formatted_document.docx"
        doc = Document()
        
        # Add a title with formatting
        title = doc.add_heading('Complex Formatted Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a subtitle
        subtitle = doc.add_heading('For Testing Document Viewer Functionality', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a paragraph with mixed formatting
        p = doc.add_paragraph('This document contains various formatting elements to test the document viewer functionality. ')
        p.add_run('Bold text, ').bold = True
        p.add_run('italic text, ').italic = True
        run = p.add_run('colored text, ')
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)  # Blue color
        p.add_run('and ').italic = True
        run = p.add_run('bold-italic-colored text.')
        run.bold = True
        run.italic = True
        run.font.color.rgb = RGBColor(0xE9, 0x24, 0x24)  # Red color
        
        # Add a paragraph with different font size
        p = doc.add_paragraph()
        run = p.add_run('This text has a different font size.')
        run.font.size = Pt(16)
        
        # Add a bulleted list
        doc.add_paragraph('Here is a bulleted list:', style='List Bullet')
        doc.add_paragraph('First bullet point with some details', style='List Bullet')
        p = doc.add_paragraph('Second bullet point with ', style='List Bullet')
        p.add_run('bold text').bold = True
        doc.add_paragraph('Third bullet point with regular text', style='List Bullet')
        
        # Add a numbered list
        doc.add_paragraph('Here is a numbered list:', style='List Number')
        doc.add_paragraph('First numbered item with some details', style='List Number')
        p = doc.add_paragraph('Second numbered item with ', style='List Number')
        p.add_run('italic text').italic = True
        doc.add_paragraph('Third numbered item with regular text', style='List Number')
        
        # Add a table
        doc.add_paragraph('Here is a table:')
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Header 1'
        header_cells[1].text = 'Header 2'
        header_cells[2].text = 'Header 3'
        
        # Add data rows
        data_rows = [
            ['Row 1, Cell 1', 'Row 1, Cell 2', 'Row 1, Cell 3'],
            ['Row 2, Cell 1', 'Row 2, Cell 2', 'Row 2, Cell 3'],
            ['Row 3, Cell 1', 'Row 3, Cell 2', 'Row 3, Cell 3']
        ]
        
        for i, row_data in enumerate(data_rows):
            row = table.rows[i+1]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text
        
        # Make some cells bold
        table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[2].cells[1].paragraphs[0].runs[0].bold = True
        table.rows[3].cells[2].paragraphs[0].runs[0].bold = True
        
        # Add a section break
        doc.add_page_break()
        
        # Add content to the second page
        doc.add_heading('Second Page Content', 1)
        doc.add_paragraph('This is the second page of our test document.')
        
        # Add another formatted paragraph
        p = doc.add_paragraph('This paragraph has ')
        p.add_run('mixed ').bold = True
        p.add_run('formatting ').italic = True
        run = p.add_run('styles ')
        run.bold = True
        run.italic = True
        p.add_run('to test preservation of complex formatting.')
        
        # Save the document
        doc.save(docx_path)
        print("✅ Complex DOCX file created successfully")
        
        # 3. Upload the complex DOCX file
        print("\n3. Uploading complex DOCX file...")
        
        with open(docx_path, "rb") as f:
            files = {"file": ("complex_formatted_document.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                headers={"Authorization": headers["Authorization"]}
            )
        
        assert response.status_code == 200, f"Failed to upload DOCX file: {response.text}"
        data = response.json()
        
        # Verify response structure
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        
        # Save document ID for retrieval test
        document_id = data["document"]["id"]
        
        # Print some details about the document
        print(f"Document ID: {document_id}")
        print(f"Document title: {data['document']['title']}")
        print(f"Total pages: {data['document']['total_pages']}")
        print(f"Contains formatting: {data['document']['metadata'].get('contains_formatting', False)}")
        
        # 4. Retrieve the document to check formatting preservation
        print("\n4. Retrieving document to check formatting preservation...")
        
        response = requests.get(f"{base_url}/documents/{document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve document: {response.text}"
        data = response.json()
        
        # Verify document structure
        assert data["id"] == document_id, "Document ID mismatch"
        assert "pages" in data, "Missing pages in retrieved document"
        assert len(data["pages"]) > 0, "No pages found in retrieved document"
        
        # Check that content is HTML
        page_content = data["pages"][0]["content"]
        assert "<div" in page_content, "Retrieved content is not in HTML format"
        
        # Check for HTML formatting elements
        formatting_elements = [
            "<b>", "<strong>",  # Bold
            "<i>", "<em>",      # Italic
            "<table>", "<tr>", "<td>",  # Table
            "<ul>", "<ol>", "<li>",     # Lists
            "<h1>", "<h2>", "<h3>",     # Headings
            "style="                     # Inline styles
        ]
        
        found_elements = []
        for element in formatting_elements:
            if element in page_content.lower():
                found_elements.append(element)
        
        print(f"Found formatting elements: {', '.join(found_elements)}")
        
        # Check if multiple pages were preserved
        assert len(data["pages"]) >= 2, "Multiple pages were not preserved"
        
        # Check metadata
        assert "metadata" in data, "Missing metadata in document"
        assert data["metadata"].get("contains_formatting", False), "contains_formatting flag should be True"
        
        # 5. Check page structure and content
        print("\n5. Checking page structure and content...")
        
        # Verify first page
        page1 = data["pages"][0]
        assert page1["page_number"] == 1, "First page should have page_number 1"
        assert "Complex Formatted Document" in page1["content"], "First page content missing title"
        assert "<table" in page1["content"].lower(), "First page content missing table"
        
        # Verify second page if it exists
        if len(data["pages"]) > 1:
            page2 = data["pages"][1]
            assert page2["page_number"] == 2, "Second page should have page_number 2"
            assert "Second Page Content" in page2["content"], "Second page content missing title"
        
        print("✅ Page structure and content verification successful")
        
        # 6. Test updating a page with complex HTML content
        print("\n6. Testing page update with complex HTML content...")
        
        complex_html = """
        <div style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <h1 style="color: #2c3e50; text-align: center;">Updated Page Content</h1>
            <p>This is a paragraph with <strong>bold text</strong>, <em>italic text</em>, and <span style="color: blue;">colored text</span>.</p>
            
            <h2>Formatted Lists</h2>
            <ul>
                <li>Unordered list item 1</li>
                <li>Unordered list item 2 with <strong>bold text</strong></li>
                <li>Unordered list item 3</li>
            </ul>
            
            <ol>
                <li>Ordered list item 1</li>
                <li>Ordered list item 2 with <em>italic text</em></li>
                <li>Ordered list item 3</li>
            </ol>
            
            <h2>Formatted Table</h2>
            <table style="width: 100%; border-collapse: collapse; margin: 20px 0;">
                <tr style="background-color: #f2f2f2;">
                    <th style="border: 1px solid #ddd; padding: 8px; text-align: left;">Header 1</th>
                    <th style="border: 1px solid #ddd; padding: 8px; text-align: left;">Header 2</th>
                    <th style="border: 1px solid #ddd; padding: 8px; text-align: left;">Header 3</th>
                </tr>
                <tr>
                    <td style="border: 1px solid #ddd; padding: 8px;"><strong>Bold Cell</strong></td>
                    <td style="border: 1px solid #ddd; padding: 8px;">Regular Cell</td>
                    <td style="border: 1px solid #ddd; padding: 8px;">Regular Cell</td>
                </tr>
                <tr>
                    <td style="border: 1px solid #ddd; padding: 8px;">Regular Cell</td>
                    <td style="border: 1px solid #ddd; padding: 8px;"><em>Italic Cell</em></td>
                    <td style="border: 1px solid #ddd; padding: 8px;">Regular Cell</td>
                </tr>
            </table>
        </div>
        """
        
        page_update_data = {
            "title": "Updated Complex Page",
            "content": complex_html
        }
        
        response = requests.put(
            f"{base_url}/documents/{document_id}/pages/1",
            json=page_update_data,
            headers=headers
        )
        
        assert response.status_code == 200, f"Failed to update document page: {response.text}"
        
        # Retrieve the document to verify the update
        response = requests.get(f"{base_url}/documents/{document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve updated document: {response.text}"
        data = response.json()
        
        # Check that the page was updated with complex HTML
        updated_page = data["pages"][0]
        assert updated_page["title"] == "Updated Complex Page", "Page title not updated"
        assert "Updated Page Content" in updated_page["content"], "Page content not updated"
        assert "<table" in updated_page["content"], "Table formatting not preserved"
        assert "<ul>" in updated_page["content"], "List formatting not preserved"
        
        print("✅ Page update with complex HTML content successful")
        
        # 7. Clean up
        print("\n7. Cleaning up test document...")
        
        # Delete the test document
        response = requests.delete(f"{base_url}/documents/{document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to delete document: {response.text}"
        
        # Clean up the test file
        if os.path.exists(docx_path):
            os.remove(docx_path)
        
        print("✅ Test cleanup successful")
        
        print("\n✅ All document viewer functionality tests passed successfully!")
        
        # Return a summary of findings
        return {
            "success": True,
            "formatting_preserved": True,
            "found_elements": found_elements,
            "multiple_pages_preserved": len(data["pages"]) >= 2,
            "contains_formatting_flag": data["metadata"].get("contains_formatting", False)
        }
        
    except ImportError as e:
        print(f"⚠️ Required library not available: {e}")
        print("To run these tests, install the required libraries: pip install python-docx Pillow")
        return {
            "success": False,
            "error": str(e)
        }
    except Exception as e:
        print(f"❌ Test failed: {e}")
        return {
            "success": False,
            "error": str(e)
        }

if __name__ == "__main__":
    results = test_document_viewer_functionality()
    print("\nTest Results Summary:")
    for key, value in results.items():
        print(f"  {key}: {value}")