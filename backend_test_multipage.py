import requests
import json
import unittest
import uuid
import os
import time
from datetime import datetime

# Get the backend URL from the frontend .env file
BACKEND_URL = "https://7056c40d-c306-42f0-bfc6-379013315fc8.preview.emergentagent.com/api"

def test_multipage_document_functionality():
    """Test the multi-page document functionality"""
    print("Starting multi-page document functionality tests...")
    
    # Test variables
    base_url = BACKEND_URL
    headers = {"Content-Type": "application/json"}
    test_user_email = f"test.user.{uuid.uuid4()}@example.com"
    test_user_password = "SecurePassword123!"
    
    # 1. Register a test user
    print("\n1. Registering test user...")
    user_data = {
        "email": test_user_email,
        "full_name": "Test User",
        "role": "editor",
        "organization": "Test Organization",
        "password": test_user_password
    }
    
    response = requests.post(f"{base_url}/auth/register", json=user_data)
    assert response.status_code == 200, f"Failed to register test user: {response.text}"
    data = response.json()
    
    # Save token for subsequent tests
    access_token = data["access_token"]
    headers["Authorization"] = f"Bearer {access_token}"
    print("✅ User registration successful")
    
    # 2. Create test files
    print("\n2. Creating test files...")
    
    # Create a multi-page PDF file
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = "/tmp/test_multipage.pdf"
        c = canvas.Canvas(pdf_path, pagesize=letter)
        
        # Add multiple pages to the PDF
        for i in range(1, 6):  # Create a 5-page PDF
            c.drawString(100, 750, f"This is page {i} of the test PDF document")
            c.drawString(100, 730, f"Page {i} content for testing multi-page functionality")
            c.drawString(100, 710, f"The document platform should split this into separate pages")
            
            # Add some more content to make the page more substantial
            for j in range(1, 10):
                c.drawString(100, 700 - j*20, f"Line {j} of additional content on page {i}")
            
            c.showPage()  # Move to the next page
        
        c.save()
        print("✅ Multi-page PDF created successfully")
        
        # Create a multi-page DOCX file
        from docx import Document
        
        docx_path = "/tmp/test_multipage.docx"
        doc = Document()
        
        # Add multiple pages to the DOCX with page breaks
        for i in range(1, 6):  # Create a 5-page DOCX
            doc.add_heading(f'Page {i} of Test Document', 0)
            doc.add_paragraph(f'This is page {i} of the test DOCX document.')
            doc.add_paragraph(f'Page {i} content for testing multi-page functionality.')
            doc.add_paragraph(f'The document platform should split this into separate pages.')
            
            # Add some more content to make the page more substantial
            for j in range(1, 10):
                doc.add_paragraph(f'Line {j} of additional content on page {i}.')
            
            # Add a page break (except for the last page)
            if i < 5:
                doc.add_page_break()
        
        doc.save(docx_path)
        print("✅ Multi-page DOCX created successfully")
        
        # 3. Test PDF upload with multi-page functionality
        print("\n3. Testing PDF upload with multi-page functionality...")
        
        # Upload the PDF file
        with open(pdf_path, "rb") as f:
            files = {"file": ("test_multipage.pdf", f, "application/pdf")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                data={"extract_text": "true"},  # Extract text for editing
                headers={"Authorization": headers["Authorization"]}
            )
        
        assert response.status_code == 200, f"Failed to upload PDF file: {response.text}"
        data = response.json()
        
        # Verify response structure
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        
        # Verify multi-page structure
        assert "pages" in data["document"], "Missing pages in document"
        assert len(data["document"]["pages"]) > 1, f"Expected multiple pages, got {len(data['document']['pages'])}"
        assert data["document"]["total_pages"] > 1, f"Expected total_pages > 1, got {data['document']['total_pages']}"
        
        # Check that each page has the correct page number
        for i, page in enumerate(data["document"]["pages"], 1):
            assert page["page_number"] == i, f"Expected page_number {i}, got {page['page_number']}"
            assert f"Page {i}" in page["title"], f"Expected 'Page {i}' in title, got {page['title']}"
        
        # Save document ID for retrieval test
        pdf_document_id = data["document"]["id"]
        
        # Print some details about the document
        print(f"PDF Document ID: {pdf_document_id}")
        print(f"PDF Document title: {data['document']['title']}")
        print(f"PDF Total pages: {data['document']['total_pages']}")
        
        print("✅ PDF upload with multi-page functionality working")
        
        # 4. Test DOCX upload with multi-page functionality
        print("\n4. Testing DOCX upload with multi-page functionality...")
        
        # Upload the DOCX file
        with open(docx_path, "rb") as f:
            files = {"file": ("test_multipage.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                headers={"Authorization": headers["Authorization"]}
            )
        
        assert response.status_code == 200, f"Failed to upload DOCX file: {response.text}"
        data = response.json()
        
        # Verify response structure
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        
        # Verify multi-page structure
        assert "pages" in data["document"], "Missing pages in document"
        assert len(data["document"]["pages"]) > 1, f"Expected multiple pages, got {len(data['document']['pages'])}"
        assert data["document"]["total_pages"] > 1, f"Expected total_pages > 1, got {data['document']['total_pages']}"
        
        # Check that each page has the correct page number
        for i, page in enumerate(data["document"]["pages"], 1):
            assert page["page_number"] == i, f"Expected page_number {i}, got {page['page_number']}"
        
        # Save document ID for retrieval test
        docx_document_id = data["document"]["id"]
        
        # Print some details about the document
        print(f"DOCX Document ID: {docx_document_id}")
        print(f"DOCX Document title: {data['document']['title']}")
        print(f"DOCX Total pages: {data['document']['total_pages']}")
        
        print("✅ DOCX upload with multi-page functionality working")
        
        # 5. Test document retrieval with multi-page structure
        print("\n5. Testing document retrieval with multi-page structure...")
        
        # Give the server a moment to process the document
        time.sleep(1)
        
        # Retrieve the PDF document
        response = requests.get(f"{base_url}/documents/{pdf_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve PDF document: {response.text}"
        data = response.json()
        
        # Verify document structure
        assert data["id"] == pdf_document_id, "Document ID mismatch"
        assert "pages" in data, "Missing pages in retrieved document"
        assert len(data["pages"]) > 1, f"Expected multiple pages, got {len(data['pages'])}"
        assert data["total_pages"] > 1, f"Expected total_pages > 1, got {data['total_pages']}"
        
        # Check that each page has the correct content
        for i, page in enumerate(data["pages"], 1):
            assert page["page_number"] == i, f"Expected page_number {i}, got {page['page_number']}"
            assert f"Page {i}" in page["title"], f"Expected 'Page {i}' in title, got {page['title']}"
            # Check that the page content is not empty
            assert page["content"], f"Page {i} content is empty"
        
        # Print some details about the retrieved document
        print(f"Retrieved PDF document title: {data['title']}")
        print(f"Retrieved PDF document total pages: {data['total_pages']}")
        
        # Retrieve the DOCX document
        response = requests.get(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve DOCX document: {response.text}"
        data = response.json()
        
        # Verify document structure
        assert data["id"] == docx_document_id, "Document ID mismatch"
        assert "pages" in data, "Missing pages in retrieved document"
        assert len(data["pages"]) > 1, f"Expected multiple pages, got {len(data['pages'])}"
        assert data["total_pages"] > 1, f"Expected total_pages > 1, got {data['total_pages']}"
        
        # Check that each page has the correct content
        for i, page in enumerate(data["pages"], 1):
            assert page["page_number"] == i, f"Expected page_number {i}, got {page['page_number']}"
            # Check that the page content is not empty
            assert page["content"], f"Page {i} content is empty"
            # Check that the page content contains the expected text
            assert f"Page {i}" in page["content"], f"Expected 'Page {i}' in content, got {page['content'][:100]}..."
        
        # Print some details about the retrieved document
        print(f"Retrieved DOCX document title: {data['title']}")
        print(f"Retrieved DOCX document total pages: {data['total_pages']}")
        
        print("✅ Document retrieval with multi-page structure working")
        
        # 6. Test updating a specific page
        print("\n6. Testing updating a specific page...")
        
        # Update page 2 of the PDF document
        page_update_data = {
            "title": "Updated Page 2 Title",
            "content": "<div style='font-family: Arial; color: blue;'>This is updated content for page 2 with <b>HTML formatting</b> and <i>styling</i>.</div>"
        }
        
        response = requests.put(
            f"{base_url}/documents/{pdf_document_id}/pages/2",
            json=page_update_data,
            headers=headers
        )
        
        assert response.status_code == 200, f"Failed to update document page: {response.text}"
        data = response.json()
        
        # Verify update response
        assert data["message"] == "Page updated successfully", "Unexpected response message"
        
        # Retrieve the document to verify the update
        response = requests.get(f"{base_url}/documents/{pdf_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve updated document: {response.text}"
        data = response.json()
        
        # Check that the page was updated
        page_2 = next((page for page in data["pages"] if page["page_number"] == 2), None)
        assert page_2, "Page 2 not found in document"
        assert page_2["title"] == "Updated Page 2 Title", f"Expected 'Updated Page 2 Title', got {page_2['title']}"
        assert "<b>HTML formatting</b>" in page_2["content"], "HTML formatting not preserved in update"
        
        print("✅ Updating a specific page working")
        
        # 7. Test adding multimedia to a specific page
        print("\n7. Testing adding multimedia to a specific page...")
        
        # Add a multimedia element to page 3 of the DOCX document
        multimedia_data = {
            "type": "image",
            "url": "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==",
            "title": "Test Image",
            "description": "A test base64 image"
        }
        
        response = requests.post(
            f"{base_url}/documents/{docx_document_id}/pages/3/multimedia",
            json=multimedia_data,
            headers=headers
        )
        
        assert response.status_code == 200, f"Failed to add multimedia to page: {response.text}"
        data = response.json()
        
        # Verify multimedia response
        assert data["message"] == "Multimedia element added to page successfully", "Unexpected response message"
        assert "element" in data, "Missing element in response"
        assert data["element"]["type"] == "image", "Element type mismatch"
        
        # Retrieve the document to verify the multimedia element
        response = requests.get(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve document with multimedia: {response.text}"
        data = response.json()
        
        # Check that the multimedia element was added to the correct page
        page_3 = next((page for page in data["pages"] if page["page_number"] == 3), None)
        assert page_3, "Page 3 not found in document"
        assert len(page_3["multimedia_elements"]) > 0, "No multimedia elements found on page 3"
        assert page_3["multimedia_elements"][0]["type"] == "image", "Multimedia element type mismatch"
        
        print("✅ Adding multimedia to a specific page working")
        
        # 8. Test adding interactive element to a specific page
        print("\n8. Testing adding interactive element to a specific page...")
        
        # Add an interactive element to page 4 of the DOCX document
        interactive_data = {
            "type": "signature_field",
            "label": "Client Signature",
            "required": True,
            "position": {"x": 0.5, "y": 0.8}
        }
        
        response = requests.post(
            f"{base_url}/documents/{docx_document_id}/pages/4/interactive",
            json=interactive_data,
            headers=headers
        )
        
        assert response.status_code == 200, f"Failed to add interactive element to page: {response.text}"
        data = response.json()
        
        # Verify interactive element response
        assert data["message"] == "Interactive element added to page successfully", "Unexpected response message"
        assert "element" in data, "Missing element in response"
        assert data["element"]["type"] == "signature_field", "Element type mismatch"
        
        # Retrieve the document to verify the interactive element
        response = requests.get(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve document with interactive element: {response.text}"
        data = response.json()
        
        # Check that the interactive element was added to the correct page
        page_4 = next((page for page in data["pages"] if page["page_number"] == 4), None)
        assert page_4, "Page 4 not found in document"
        assert len(page_4["interactive_elements"]) > 0, "No interactive elements found on page 4"
        assert page_4["interactive_elements"][0]["type"] == "signature_field", "Interactive element type mismatch"
        
        print("✅ Adding interactive element to a specific page working")
        
        # 9. Test page view tracking
        print("\n9. Testing page view tracking...")
        
        # Track a view of page 1 of the PDF document
        view_data = {
            "document_id": pdf_document_id,
            "page_number": 1,
            "viewer_info": {
                "ip_address": "192.168.1.1",
                "user_agent": "Test User Agent"
            },
            "time_spent": 30,
            "scroll_depth": 0.8,
            "interactions": ["click", "scroll"]
        }
        
        response = requests.post(f"{base_url}/tracking/page-view", json=view_data)
        assert response.status_code == 200, f"Failed to track page view: {response.text}"
        data = response.json()
        
        # Verify response
        assert data["message"] == "Page view tracked successfully", f"Expected 'Page view tracked successfully', got {data['message']}"
        assert "session_id" in data, "Missing session_id in response"
        
        # Save session ID for subsequent tracking
        session_id = data["session_id"]
        
        # Track a view of page 2 of the PDF document with the same session
        view_data = {
            "document_id": pdf_document_id,
            "page_number": 2,
            "session_id": session_id,
            "viewer_info": {
                "ip_address": "192.168.1.1",
                "user_agent": "Test User Agent"
            },
            "time_spent": 45,
            "scroll_depth": 1.0,
            "interactions": ["click", "scroll", "highlight"]
        }
        
        response = requests.post(f"{base_url}/tracking/page-view", json=view_data)
        assert response.status_code == 200, f"Failed to track page view: {response.text}"
        
        print("✅ Page view tracking working")
        
        # 10. Test page analytics
        print("\n10. Testing page analytics...")
        
        # Get page analytics for the PDF document
        response = requests.get(f"{base_url}/documents/{pdf_document_id}/page-analytics", headers=headers)
        assert response.status_code == 200, f"Failed to get page analytics: {response.text}"
        data = response.json()
        
        # Verify analytics structure
        assert "document_id" in data, "Missing document_id in response"
        assert "total_pages" in data, "Missing total_pages in response"
        assert "page_analytics" in data, "Missing page_analytics in response"
        assert "overall_analytics" in data, "Missing overall_analytics in response"
        
        # Check that analytics for each page are present
        assert len(data["page_analytics"]) == data["total_pages"], f"Expected {data['total_pages']} page analytics, got {len(data['page_analytics'])}"
        
        # Check that page 1 and 2 have view data
        page_1_analytics = next((page for page in data["page_analytics"] if page["page_number"] == 1), None)
        page_2_analytics = next((page for page in data["page_analytics"] if page["page_number"] == 2), None)
        
        assert page_1_analytics, "Page 1 analytics not found"
        assert page_2_analytics, "Page 2 analytics not found"
        
        assert page_1_analytics["total_views"] > 0, f"Expected page 1 views > 0, got {page_1_analytics['total_views']}"
        assert page_2_analytics["total_views"] > 0, f"Expected page 2 views > 0, got {page_2_analytics['total_views']}"
        
        print("✅ Page analytics working")
        
        # 11. Clean up
        print("\n11. Cleaning up test documents...")
        
        # Delete the test documents
        response = requests.delete(f"{base_url}/documents/{pdf_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to delete PDF document: {response.text}"
        
        response = requests.delete(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to delete DOCX document: {response.text}"
        
        # Clean up the test files
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        
        if os.path.exists(docx_path):
            os.remove(docx_path)
        
        print("✅ Test cleanup successful")
        
        print("\n✅ All multi-page document functionality tests passed successfully!")
    
    except ImportError as e:
        print(f"⚠️ Required library not available: {e}")
        print("To run these tests, install the required libraries: pip install reportlab python-docx")
    except Exception as e:
        print(f"❌ Test failed: {e}")
        # Clean up test files if they exist
        for path in ["/tmp/test_multipage.pdf", "/tmp/test_multipage.docx"]:
            if os.path.exists(path):
                os.remove(path)
        raise

if __name__ == "__main__":
    test_multipage_document_functionality()