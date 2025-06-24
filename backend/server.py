import sys
from pathlib import Path

# Add the backend directory to Python path
backend_dir = Path(__file__).parent
sys.path.append(str(backend_dir))

from fastapi import FastAPI, APIRouter, Depends, HTTPException, status, File, UploadFile, Form
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
import os
import logging
from typing import List, Optional, Dict, Any
from datetime import datetime, timedelta
import json

# Import our custom modules
from models import *
from database import Database, get_collection, create_indexes
from auth import get_current_active_user, require_role, check_document_access, create_access_token, verify_password, get_password_hash
from openai_service import openai_service

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Create the main app
app = FastAPI(
    title="Document Platform API",
    description="AI-powered document creation and collaboration platform",
    version="1.0.0"
)

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ==================== AUTH ENDPOINTS ====================

@api_router.post("/auth/register")
async def register_user(user_data: UserCreate):
    """Register a new user"""
    users_collection = await get_collection('users')
    
    # Check if user already exists
    existing_user = await users_collection.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create new user
    hashed_password = get_password_hash(user_data.password)
    user = User(
        email=user_data.email,
        full_name=user_data.full_name,
        role=user_data.role,
        organization=user_data.organization
    )
    
    # Save user (password stored separately for security)
    user_dict = user.dict()
    user_dict['hashed_password'] = hashed_password
    
    await users_collection.insert_one(user_dict)
    
    # Create access token
    access_token = create_access_token(data={"sub": user.id})
    
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": user
    }

@api_router.post("/auth/login")
async def login_user(login_data: dict):
    """Authenticate user and return access token"""
    users_collection = await get_collection('users')
    
    email = login_data.get("email")
    password = login_data.get("password")
    
    if not email or not password:
        raise HTTPException(status_code=400, detail="Email and password are required")
    
    user_data = await users_collection.find_one({"email": email})
    if not user_data or not verify_password(password, user_data['hashed_password']):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Update last login
    await users_collection.update_one(
        {"email": email},
        {"$set": {"last_login": datetime.utcnow()}}
    )
    
    user = User(**{k: v for k, v in user_data.items() if k != 'hashed_password'})
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.id}, 
        expires_delta=access_token_expires
    )
    
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": user
    }

# ==================== USER ENDPOINTS ====================

@api_router.get("/users/me", response_model=User)
async def get_current_user_profile(current_user: User = Depends(get_current_active_user)):
    """Get current user profile"""
    return current_user

@api_router.put("/users/me", response_model=User)
async def update_current_user(
    user_update: UserUpdate,
    current_user: User = Depends(get_current_active_user)
):
    """Update current user profile"""
    users_collection = await get_collection('users')
    
    update_data = {k: v for k, v in user_update.dict().items() if v is not None}
    if update_data:
        await users_collection.update_one(
            {"id": current_user.id},
            {"$set": update_data}
        )
    
    # Return updated user
    updated_user_data = await users_collection.find_one({"id": current_user.id})
    return User(**{k: v for k, v in updated_user_data.items() if k != 'hashed_password'})

# ==================== EMAIL INTEGRATION ENDPOINTS ====================

@api_router.post("/users/me/email-connections")
async def add_email_connection(
    connection_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user)
):
    """Add email connection to user profile"""
    users_collection = await get_collection('users')
    
    # Create new email connection
    email_connection = EmailConnection(
        provider=connection_data["provider"],
        email_address=connection_data["email_address"],
        display_name=connection_data.get("display_name", connection_data["email_address"]),
        is_primary=connection_data.get("is_primary", False)
    )
    
    # If this is set as primary, make others non-primary
    if email_connection.is_primary:
        await users_collection.update_one(
            {"id": current_user.id},
            {"$set": {"email_connections.$[].is_primary": False}}
        )
    
    # Add connection to user
    await users_collection.update_one(
        {"id": current_user.id},
        {"$push": {"email_connections": email_connection.dict()}}
    )
    
    return email_connection

@api_router.get("/users/me/email-connections")
async def get_email_connections(current_user: User = Depends(get_current_active_user)):
    """Get user's email connections"""
    users_collection = await get_collection('users')
    user_data = await users_collection.find_one({"id": current_user.id})
    
    return user_data.get("email_connections", [])

@api_router.delete("/users/me/email-connections/{connection_id}")
async def remove_email_connection(
    connection_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Remove email connection"""
    users_collection = await get_collection('users')
    
    # Remove the connection
    await users_collection.update_one(
        {"id": current_user.id},
        {"$pull": {"email_connections": {"id": connection_id}}}
    )
    
    return {"message": "Email connection removed successfully"}

@api_router.put("/users/me/email-connections/{connection_id}/primary")
async def set_primary_email_connection(
    connection_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Set email connection as primary"""
    users_collection = await get_collection('users')
    
    # First, make all connections non-primary
    await users_collection.update_one(
        {"id": current_user.id},
        {"$set": {"email_connections.$[].is_primary": False}}
    )
    
    # Then set the specified connection as primary
    await users_collection.update_one(
        {"id": current_user.id, "email_connections.id": connection_id},
        {"$set": {"email_connections.$.is_primary": True}}
    )
    
    return {"message": "Primary email connection updated"}

@api_router.put("/users/me/notification-settings")
async def update_notification_settings(
    settings: Dict[str, bool],
    current_user: User = Depends(get_current_active_user)
):
    """Update user notification settings"""
    users_collection = await get_collection('users')
    
    await users_collection.update_one(
        {"id": current_user.id},
        {"$set": {"notification_settings": settings}}
    )
    
    return {"message": "Notification settings updated successfully"}

@api_router.put("/users/me/email-signature")
async def update_email_signature(
    signature_data: Dict[str, str],
    current_user: User = Depends(get_current_active_user)
):
    """Update user email signature"""
    users_collection = await get_collection('users')
    
    await users_collection.update_one(
        {"id": current_user.id},
        {"$set": {"email_signature": signature_data["signature"]}}
    )
    
    return {"message": "Email signature updated successfully"}

# ==================== EMAIL SENDING ENDPOINTS ====================

@api_router.post("/documents/{document_id}/send-email")
async def send_document_via_email(
    document_id: str,
    email_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user)
):
    """Send document via email"""
    document_data = await check_document_access(document_id, current_user, "view")
    
    # Here you would integrate with email service (SMTP, SendGrid, etc.)
    # For now, we'll log the action and return success
    
    # Log the email send activity
    await log_activity(
        current_user.id, 
        current_user.full_name, 
        document_id, 
        ActionType.CREATE,  # Using CREATE for email send action
        {
            "action": "email_sent",
            "recipients": email_data.get("recipients", []),
            "subject": email_data.get("subject", ""),
            "message": email_data.get("message", "")
        }
    )
    
    return {
        "message": "Email sent successfully",
        "recipients": email_data.get("recipients", []),
        "document_title": document_data["title"]
    }

# ==================== DOCUMENT ENDPOINTS ====================

@api_router.post("/documents", response_model=Document)
async def create_document(
    document_data: DocumentCreate,
    current_user: User = Depends(get_current_active_user)
):
    """Create a new document"""
    documents_collection = await get_collection('documents')
    
    document = Document(
        title=document_data.title,
        type=document_data.type,
        owner_id=current_user.id,
        organization=current_user.organization,
        sections=document_data.sections,
        collaborators=document_data.collaborators,
        tags=document_data.tags,
        metadata=document_data.metadata
    )
    
    await documents_collection.insert_one(document.dict())
    
    # Log activity
    await log_activity(current_user.id, current_user.full_name, document.id, ActionType.CREATE, {"title": document.title})
    
    return document

@api_router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    title: str = None,
    extract_text: bool = Form(True),  # Option to extract text or preserve original
    current_user: User = Depends(get_current_active_user)
):
    """Upload a document file and convert it to Google Docs-like format with formatting preserved"""
    import os
    import base64
    from io import BytesIO
    
    try:
        import mammoth
        import docx
        from PIL import Image
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"Required library not available: {str(e)}")
    
    # Validate file type
    allowed_types = [
        "application/pdf", 
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        "text/plain"
    ]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, DOCX, or TXT files.")
    
    try:
        # Read file content
        content = await file.read()
        
        # Process document based on type to preserve formatting
        pages = []
        document_html = ""
        
        if file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                # Use Mammoth to convert DOCX to HTML with formatting preserved
                result = mammoth.convert_to_html(BytesIO(content))
                document_html = result.value  # The generated HTML
                
                # Handle images and convert to base64
                def convert_image_to_base64(image):
                    try:
                        image_bytes = image.open()
                        pil_image = Image.open(image_bytes)
                        buffer = BytesIO()
                        # Convert to RGB if necessary
                        if pil_image.mode in ('RGBA', 'LA', 'P'):
                            pil_image = pil_image.convert('RGB')
                        pil_image.save(buffer, format='JPEG', quality=90)
                        img_data = buffer.getvalue()
                        buffer.close()
                        return f"data:image/jpeg;base64,{base64.b64encode(img_data).decode()}"
                    except Exception as e:
                        print(f"Image conversion error: {e}")
                        return ""
                
                # Convert images to base64 and embed them
                if hasattr(result, 'images'):
                    for image in result.images:
                        base64_image = convert_image_to_base64(image)
                        if base64_image:
                            # Replace image references with base64 data
                            document_html = document_html.replace(
                                f'src="{image.alt_text}"', 
                                f'src="{base64_image}"'
                            )
                
                # Clean up and enhance HTML styling - PRESERVE ORIGINAL FORMATTING
                document_html = f"""
                <div style="font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
                           line-height: 1.6; 
                           width: 100%;
                           text-align: left;
                           padding: 20px;
                           background: white;">
                    {document_html}
                </div>
                """
                
                # Split into logical pages based on page breaks or content length
                # For now, create one page with all content
                page_obj = DocumentPage(
                    page_number=1,
                    title=title or file.filename.rsplit('.', 1)[0],
                    content=document_html  # Store as HTML instead of plain text
                )
                pages.append(page_obj)
                
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to process DOCX with formatting: {str(e)}")
        
        elif file.content_type == "application/pdf":
            try:
                if not extract_text:
                    # Store PDF as base64 for native PDF viewing (preserves layout, images, etc.)
                    pdf_base64 = base64.b64encode(content).decode()
                    page_obj = DocumentPage(
                        page_number=1,
                        title=title or file.filename.rsplit('.', 1)[0],
                        content=f"""
                        <div style="width: 100%; height: 800px; text-align: center;">
                            <embed src="data:application/pdf;base64,{pdf_base64}" 
                                   type="application/pdf" 
                                   width="100%" 
                                   height="100%" 
                                   style="border: none; box-shadow: 0 4px 8px rgba(0,0,0,0.1);">
                            <p style="margin-top: 20px;">
                                <a href="data:application/pdf;base64,{pdf_base64}" 
                                   download="{file.filename}" 
                                   style="color: #1a73e8; text-decoration: none;">
                                    📄 Download Original PDF
                                </a>
                            </p>
                        </div>
                        """,
                        metadata={"is_pdf": True, "original_filename": file.filename}
                    )
                    pages.append(page_obj)
                else:
                    # Extract text for editing (current behavior)
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                # Convert plain text to properly formatted HTML - PRESERVE ALIGNMENT
                                formatted_html = f"""
                                <div style="font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
                                           line-height: 1.6; 
                                           white-space: pre-wrap; 
                                           text-align: left;
                                           width: 100%;
                                           padding: 20px;">
                                    {page_text.strip()}
                                </div>
                                """
                                page_obj = DocumentPage(
                                    page_number=page_num,
                                    title=f"Page {page_num}",
                                    content=formatted_html
                                )
                                pages.append(page_obj)
                        except Exception as e:
                            page_obj = DocumentPage(
                                page_number=page_num,
                                title=f"Page {page_num}",
                                content=f"<p>[Content extraction failed: {str(e)}]</p>"
                            )
                            pages.append(page_obj)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to process PDF: {str(e)}")
        
        elif file.content_type == "text/plain":
            try:
                # Plain text - convert to HTML with proper formatting - PRESERVE ALIGNMENT
                text_content = content.decode('utf-8')
                
                # Convert line breaks to proper HTML
                formatted_html = f"""
                <div style="font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
                           line-height: 1.6; 
                           white-space: pre-wrap; 
                           text-align: left;
                           width: 100%;
                           padding: 20px;">
                    {text_content}
                </div>
                """
                
                page_obj = DocumentPage(
                    page_number=1,
                    title="Document Content",
                    content=formatted_html
                )
                pages.append(page_obj)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to process text file: {str(e)}")
        
        # If no pages created, create a default page
        if not pages:
            page_obj = DocumentPage(
                page_number=1,
                title="Page 1",
                content="<p>Document content could not be extracted properly. Please edit this page.</p>"
            )
            pages.append(page_obj)
        
        # Create legacy sections for backward compatibility (convert HTML to markdown for sections)
        sections = []
        for i, page in enumerate(pages):
            # For sections, we'll store a simplified version
            section = DocumentSection(
                title=page.title,
                content=page.content,  # Keep HTML content for rich editing
                order=i + 1
            )
            sections.append(section)
        
        # Create document
        documents_collection = await get_collection('documents')
        
        document = Document(
            title=title or file.filename.rsplit('.', 1)[0],
            type=DocumentType.PROPOSAL,
            owner_id=current_user.id,
            organization=current_user.organization,
            sections=sections,  # HTML content preserved
            pages=pages,  # HTML content preserved
            total_pages=len(pages),
            collaborators=[],
            tags=["uploaded", "imported", "formatted"],
            metadata={
                "original_filename": file.filename,
                "file_type": file.content_type,
                "upload_method": "formatted_upload",
                "total_pages": len(pages),
                "processing_method": "html_preservation",
                "contains_formatting": True
            }
        )
        
        await documents_collection.insert_one(document.dict())
        
        # Log activity
        await log_activity(
            current_user.id, 
            current_user.full_name, 
            document.id, 
            ActionType.CREATE, 
            {
                "title": document.title,
                "upload_method": "file_upload",
                "original_filename": file.filename,
                "total_pages": len(pages)
            }
        )
        
        return {
            "message": "Document uploaded and processed successfully",
            "document": document,
            "total_pages": len(pages),
            "processing_method": "page_wise_extraction"
        }
        
    except Exception as e:
        logger.error(f"Error processing uploaded file: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@api_router.get("/documents", response_model=List[Document])
async def get_documents(
    type: Optional[DocumentType] = None,
    status: Optional[DocumentStatus] = None,
    current_user: User = Depends(get_current_active_user)
):
    """Get user's documents with optional filtering"""
    documents_collection = await get_collection('documents')
    
    # Build query
    query = {
        "$or": [
            {"owner_id": current_user.id},
            {"collaborators.user_id": current_user.id}
        ]
    }
    
    if type:
        query["type"] = type
    if status:
        query["status"] = status
    
    documents = await documents_collection.find(query).to_list(1000)
    
    # Clean up documents to handle None values in sections
    cleaned_documents = []
    for doc in documents:
        # Clean sections - handle None titles
        if 'sections' in doc and doc['sections']:
            for section in doc['sections']:
                if section.get('title') is None:
                    section['title'] = "Untitled Section"
                if section.get('content') is None:
                    section['content'] = ""
        
        # Clean pages - handle None titles  
        if 'pages' in doc and doc['pages']:
            for page in doc['pages']:
                if page.get('title') is None:
                    page['title'] = f"Page {page.get('page_number', 1)}"
                if page.get('content') is None:
                    page['content'] = ""
        
        cleaned_documents.append(doc)
    
    return [Document(**doc) for doc in cleaned_documents]

@api_router.get("/documents/{document_id}", response_model=Document)
async def get_document(
    document_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Get a specific document"""
    document_data = await check_document_access(document_id, current_user, "view")
    
    # Clean document data to handle None values
    if 'sections' in document_data and document_data['sections']:
        for section in document_data['sections']:
            if section.get('title') is None:
                section['title'] = "Untitled Section"
            if section.get('content') is None:
                section['content'] = ""
    
    if 'pages' in document_data and document_data['pages']:
        for page in document_data['pages']:
            if page.get('title') is None:
                page['title'] = f"Page {page.get('page_number', 1)}"
            if page.get('content') is None:
                page['content'] = ""
    
    # Log view activity
    await log_activity(current_user.id, current_user.full_name, document_id, ActionType.VIEW, {})
    
    return Document(**document_data)

@api_router.put("/documents/{document_id}", response_model=Document)
async def update_document(
    document_id: str,
    document_update: DocumentUpdate,
    current_user: User = Depends(get_current_active_user)
):
    """Update a document"""
    await check_document_access(document_id, current_user, "edit")
    documents_collection = await get_collection('documents')
    
    update_data = {k: v for k, v in document_update.dict().items() if v is not None}
    update_data["updated_at"] = datetime.utcnow()
    
    await documents_collection.update_one(
        {"id": document_id},
        {"$set": update_data}
    )
    
    # Log activity
    await log_activity(current_user.id, current_user.full_name, document_id, ActionType.EDIT, update_data)
    
    # Return updated document
    updated_doc = await documents_collection.find_one({"id": document_id})
    return Document(**updated_doc)

@api_router.put("/documents/{document_id}/sections/{section_id}")
async def update_document_section(
    document_id: str,
    section_id: str,
    section_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user)
):
    """Update a specific section of a document"""
    await check_document_access(document_id, current_user, "edit")
    documents_collection = await get_collection('documents')
    
    # Update the specific section
    await documents_collection.update_one(
        {"id": document_id, "sections.id": section_id},
        {"$set": {
            "sections.$.title": section_data.get("title"),
            "sections.$.content": section_data.get("content"),
            "sections.$.multimedia_elements": section_data.get("multimedia_elements", []),
            "sections.$.interactive_elements": section_data.get("interactive_elements", []),
            "updated_at": datetime.utcnow()
        }}
    )
    
    # Log activity
    await log_activity(
        current_user.id, 
        current_user.full_name, 
        document_id, 
        ActionType.EDIT, 
        {
            "section_id": section_id,
            "section_title": section_data.get("title"),
            "action": "section_updated"
        }
    )
    
    return {"message": "Section updated successfully"}

@api_router.post("/documents/{document_id}/sections/{section_id}/multimedia")
async def add_multimedia_element(
    document_id: str,
    section_id: str,
    multimedia_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user)
):
    """Add multimedia element to a document section"""
    await check_document_access(document_id, current_user, "edit")
    documents_collection = await get_collection('documents')
    
    multimedia_element = MultimediaElement(
        type=multimedia_data["type"],
        url=multimedia_data["url"],
        title=multimedia_data.get("title"),
        description=multimedia_data.get("description"),
        duration=multimedia_data.get("duration")
    )
    
    # Add multimedia element to the section
    await documents_collection.update_one(
        {"id": document_id, "sections.id": section_id},
        {"$push": {"sections.$.multimedia_elements": multimedia_element.dict()}}
    )
    
    # Log activity
    await log_activity(
        current_user.id, 
        current_user.full_name, 
        document_id, 
        ActionType.EDIT, 
        {
            "section_id": section_id,
            "action": "multimedia_added",
            "multimedia_type": multimedia_data["type"]
        }
    )
    
    return {"message": "Multimedia element added successfully", "element": multimedia_element}

@api_router.post("/documents/{document_id}/sections/{section_id}/interactive")
async def add_interactive_element(
    document_id: str,
    section_id: str,
    interactive_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user)
):
    """Add interactive element (button, e-signature field) to a document section"""
    await check_document_access(document_id, current_user, "edit")
    documents_collection = await get_collection('documents')
    
    interactive_element = InteractiveElement(
        type=interactive_data["type"],
        label=interactive_data["label"],
        action=interactive_data.get("action"),
        required=interactive_data.get("required", False),
        position=interactive_data.get("position", {"x": 0.5, "y": 0.3, "page": 1})
    )
    
    # Add interactive element to the section
    await documents_collection.update_one(
        {"id": document_id, "sections.id": section_id},
        {"$push": {"sections.$.interactive_elements": interactive_element.dict()}}
    )
    
    # Log activity
    await log_activity(
        current_user.id, 
        current_user.full_name, 
        document_id, 
        ActionType.EDIT, 
        {
            "section_id": section_id,
            "action": "interactive_added",
            "interactive_type": interactive_data["type"]
        }
    )
    
    return {"message": "Interactive element added successfully", "element": interactive_element}

@api_router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Delete a document"""
    await check_document_access(document_id, current_user, "admin")
    documents_collection = await get_collection('documents')
    
    result = await documents_collection.delete_one({"id": document_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {"message": "Document deleted successfully"}

# ==================== AI/RFP GENERATION ENDPOINTS ====================

@api_router.post("/ai/generate-rfp")
async def generate_rfp(
    request: RFPGenerationRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Generate RFP content using AI"""
    try:
        sections = await openai_service.generate_rfp_content(request)
        
        return {
            "sections": [section.dict() for section in sections],
            "message": "RFP content generated successfully"
        }
    except Exception as e:
        logger.error(f"Error generating RFP: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/ai/analyze-document/{document_id}")
async def analyze_document_performance(
    document_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Analyze document performance and get AI recommendations"""
    document_data = await check_document_access(document_id, current_user, "view")
    
    # Get performance data
    performance_data = await get_document_performance_data(document_id)
    
    # Generate AI recommendations
    recommendations = await openai_service.analyze_document_performance(document_data, performance_data)
    
    # Save recommendations
    if recommendations:
        recommendations_collection = await get_collection('ai_recommendations')
        for rec in recommendations:
            await recommendations_collection.insert_one(rec.dict())
    
    return {
        "recommendations": [rec.dict() for rec in recommendations],
        "performance_data": performance_data
    }

# ==================== COLLABORATION ENDPOINTS ====================

@api_router.post("/documents/{document_id}/comments")
async def add_comment(
    document_id: str,
    comment_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user)
):
    """Add a comment to a document"""
    await check_document_access(document_id, current_user, "comment")
    documents_collection = await get_collection('documents')
    
    comment = Comment(
        user_id=current_user.id,
        user_name=current_user.full_name,
        content=comment_data["content"],
        section_id=comment_data.get("section_id"),
        position=comment_data.get("position")
    )
    
    await documents_collection.update_one(
        {"id": document_id},
        {"$push": {"comments": comment.dict()}}
    )
    
    # Log activity
    await log_activity(current_user.id, current_user.full_name, document_id, ActionType.COMMENT, {"content": comment.content})
    
    return comment

@api_router.get("/documents/{document_id}/comments")
async def get_comments(
    document_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Get all comments for a document"""
    document_data = await check_document_access(document_id, current_user, "view")
    return document_data.get("comments", [])

# ==================== PAGE-WISE TRACKING ENDPOINTS ====================

@api_router.post("/tracking/page-view")
async def track_page_view(view_data: Dict[str, Any]):
    """Track page-wise document viewing with detailed analytics"""
    views_collection = await get_collection('document_views')
    
    # Check if session already exists
    session_id = view_data.get("session_id")
    document_id = view_data["document_id"]
    page_number = view_data["page_number"]
    
    if session_id:
        # Update existing session
        existing_view = await views_collection.find_one({"session_id": session_id})
        
        if existing_view:
            # Update page view data
            pages_viewed = existing_view.get("pages_viewed", [])
            
            # Find existing page view or create new one
            page_view_index = next((i for i, pv in enumerate(pages_viewed) if pv["page_number"] == page_number), None)
            
            page_view_data = {
                "page_number": page_number,
                "time_spent": view_data.get("time_spent", 0),
                "scroll_depth": view_data.get("scroll_depth", 0),
                "interactions": view_data.get("interactions", []),
                "entry_time": existing_view["timestamp"] if page_view_index is None else pages_viewed[page_view_index]["entry_time"],
                "exit_time": datetime.utcnow().isoformat(),
                "clicks": view_data.get("clicks", []),
                "focus_areas": view_data.get("focus_areas", [])
            }
            
            if page_view_index is not None:
                pages_viewed[page_view_index] = page_view_data
            else:
                pages_viewed.append(page_view_data)
            
            # Update session
            await views_collection.update_one(
                {"session_id": session_id},
                {
                    "$set": {
                        "pages_viewed": pages_viewed,
                        "current_page": page_number,
                        "max_page_reached": max(existing_view.get("max_page_reached", 1), page_number),
                        "total_time_spent": sum(pv.get("time_spent", 0) for pv in pages_viewed)
                    }
                }
            )
        else:
            # Create new session if not found
            session_id = str(uuid.uuid4())
    
    if not session_id or not existing_view:
        # Create new document view session
        session_id = str(uuid.uuid4())
        
        page_view_data = PageView(
            page_number=page_number,
            time_spent=view_data.get("time_spent", 0),
            scroll_depth=view_data.get("scroll_depth", 0),
            interactions=view_data.get("interactions", []),
            clicks=view_data.get("clicks", []),
            focus_areas=view_data.get("focus_areas", [])
        )
        
        document_view = DocumentView(
            document_id=document_id,
            session_id=session_id,
            viewer_info=ViewerInfo(**view_data["viewer_info"]),
            pages_viewed=[page_view_data.dict()],
            current_page=page_number,
            max_page_reached=page_number,
            total_time_spent=view_data.get("time_spent", 0)
        )
        
        await views_collection.insert_one(document_view.dict())
    
    return {"message": "Page view tracked successfully", "session_id": session_id}

@api_router.get("/documents/{document_id}/page-analytics")
async def get_page_analytics(
    document_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Get detailed page-wise analytics for a document"""
    await check_document_access(document_id, current_user, "view")
    
    views_collection = await get_collection('document_views')
    documents_collection = await get_collection('documents')
    
    # Get document info
    document = await documents_collection.find_one({"id": document_id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    total_pages = document.get("total_pages", len(document.get("pages", [])))
    
    # Get all views for this document
    views = await views_collection.find({"document_id": document_id}).to_list(1000)
    
    if not views:
        return {
            "document_id": document_id,
            "total_pages": total_pages,
            "page_analytics": [],
            "overall_analytics": {
                "total_views": 0,
                "unique_viewers": 0,
                "average_completion_rate": 0,
                "drop_off_points": []
            }
        }
    
    # Calculate page-wise analytics
    page_analytics = []
    unique_viewers = len(set(view.get("viewer_info", {}).get("ip_address") for view in views))
    
    for page_num in range(1, total_pages + 1):
        # Get all page views for this page
        page_views = []
        for view in views:
            for page_view in view.get("pages_viewed", []):
                if page_view.get("page_number") == page_num:
                    page_views.append(page_view)
        
        if page_views:
            total_page_views = len(page_views)
            avg_time = sum(pv.get("time_spent", 0) for pv in page_views) / total_page_views
            interaction_count = sum(len(pv.get("interactions", [])) for pv in page_views)
            
            # Calculate heat map data
            heat_map_data = []
            for pv in page_views:
                heat_map_data.extend(pv.get("clicks", []))
                heat_map_data.extend(pv.get("focus_areas", []))
            
            page_analytics.append({
                "page_number": page_num,
                "total_views": total_page_views,
                "unique_viewers": len(set(pv.get("viewer_id") for pv in page_views if pv.get("viewer_id"))),
                "average_time_spent": avg_time,
                "completion_rate": (total_page_views / len(views)) * 100,
                "interaction_rate": (interaction_count / total_page_views) * 100 if total_page_views > 0 else 0,
                "heat_map_data": heat_map_data
            })
        else:
            page_analytics.append({
                "page_number": page_num,
                "total_views": 0,
                "unique_viewers": 0,
                "average_time_spent": 0,
                "completion_rate": 0,
                "interaction_rate": 0,
                "heat_map_data": []
            })
    
    # Calculate drop-off points
    page_reach_counts = {}
    for view in views:
        max_page = view.get("max_page_reached", 1)
        for page_num in range(1, max_page + 1):
            page_reach_counts[page_num] = page_reach_counts.get(page_num, 0) + 1
    
    drop_off_points = []
    for page_num in range(1, total_pages):
        current_page_views = page_reach_counts.get(page_num, 0)
        next_page_views = page_reach_counts.get(page_num + 1, 0)
        if current_page_views > 0:
            drop_off_rate = ((current_page_views - next_page_views) / current_page_views) * 100
            if drop_off_rate > 30:  # Consider 30%+ drop-off as significant
                drop_off_points.append(page_num)
    
    # Overall completion rate
    completion_rate = (page_reach_counts.get(total_pages, 0) / len(views)) * 100 if views else 0
    
    return {
        "document_id": document_id,
        "total_pages": total_pages,
        "page_analytics": page_analytics,
        "overall_analytics": {
            "total_views": len(views),
            "unique_viewers": unique_viewers,
            "average_completion_rate": completion_rate,
            "drop_off_points": drop_off_points,
            "total_time_spent": sum(view.get("total_time_spent", 0) for view in views),
            "average_session_duration": sum(view.get("total_time_spent", 0) for view in views) / len(views) if views else 0
        }
    }

@api_router.put("/documents/{document_id}/pages/{page_number}")
async def update_document_page(
    document_id: str,
    page_number: int,
    page_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user)
):
    """Update a specific page of a document"""
    await check_document_access(document_id, current_user, "edit")
    documents_collection = await get_collection('documents')
    
    # Update the specific page
    await documents_collection.update_one(
        {"id": document_id, "pages.page_number": page_number},
        {"$set": {
            "pages.$.title": page_data.get("title"),
            "pages.$.content": page_data.get("content"),
            "pages.$.multimedia_elements": page_data.get("multimedia_elements", []),
            "pages.$.interactive_elements": page_data.get("interactive_elements", []),
            "updated_at": datetime.utcnow()
        }}
    )
    
    # Log activity
    await log_activity(
        current_user.id, 
        current_user.full_name, 
        document_id, 
        ActionType.EDIT, 
        {
            "page_number": page_number,
            "page_title": page_data.get("title"),
            "action": "page_updated"
        }
    )
    
    return {"message": "Page updated successfully"}

@api_router.post("/documents/{document_id}/pages/{page_number}/multimedia")
async def add_multimedia_to_page(
    document_id: str,
    page_number: int,
    multimedia_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user)
):
    """Add multimedia element to a specific page"""
    await check_document_access(document_id, current_user, "edit")
    documents_collection = await get_collection('documents')
    
    multimedia_element = MultimediaElement(
        type=multimedia_data["type"],
        url=multimedia_data["url"],
        title=multimedia_data.get("title"),
        description=multimedia_data.get("description"),
        duration=multimedia_data.get("duration"),
        position=multimedia_data.get("position", {"x": 0.5, "y": 0.5}),
        size=multimedia_data.get("size", {"width": "200px", "height": "150px"})
    )
    
    # Add multimedia element to the page
    await documents_collection.update_one(
        {"id": document_id, "pages.page_number": page_number},
        {"$push": {"pages.$.multimedia_elements": multimedia_element.dict()}}
    )
    
    return {"message": "Multimedia element added to page successfully", "element": multimedia_element}

@api_router.post("/documents/{document_id}/pages/{page_number}/interactive")
async def add_interactive_to_page(
    document_id: str,
    page_number: int,
    interactive_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user)
):
    """Add interactive element to a specific page"""
    await check_document_access(document_id, current_user, "edit")
    documents_collection = await get_collection('documents')
    
    interactive_element = InteractiveElement(
        type=interactive_data["type"],
        label=interactive_data["label"],
        action=interactive_data.get("action"),
        required=interactive_data.get("required", False),
        position=interactive_data.get("position", {"x": 0.5, "y": 0.3}),
        size=interactive_data.get("size", {"width": "120px", "height": "40px"})
    )
    
    # Add interactive element to the page
    await documents_collection.update_one(
        {"id": document_id, "pages.page_number": page_number},
        {"$push": {"pages.$.interactive_elements": interactive_element.dict()}}
    )
    
    return {"message": "Interactive element added to page successfully", "element": interactive_element}

@api_router.get("/documents/{document_id}/analytics")
async def get_document_analytics(
    document_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Get analytics for a document"""
    await check_document_access(document_id, current_user, "view")
    
    performance_data = await get_document_performance_data(document_id)
    return performance_data

# ==================== HELPER FUNCTIONS ====================

async def log_activity(user_id: str, user_name: str, document_id: str, action: ActionType, details: Dict[str, Any]):
    """Log user activity"""
    activity_logs_collection = await get_collection('activity_logs')
    
    activity = ActivityLog(
        user_id=user_id,
        user_name=user_name,
        document_id=document_id,
        action=action,
        details=details
    )
    
    await activity_logs_collection.insert_one(activity.dict())

async def get_document_performance_data(document_id: str) -> Dict[str, Any]:
    """Get performance analytics for a document"""
    views_collection = await get_collection('document_views')
    
    # Get all views for this document
    views = await views_collection.find({"document_id": document_id}).to_list(1000)
    
    if not views:
        return {
            "views": 0,
            "unique_viewers": 0,
            "avg_time": 0,
            "completion_rate": 0,
            "download_rate": 0,
            "sign_rate": 0,
            "engagement": {}
        }
    
    total_views = len(views)
    unique_viewers = len(set(view.get("viewer_info", {}).get("ip_address") for view in views))
    avg_time = sum(view.get("total_time_spent", 0) for view in views) / total_views
    completion_rate = (sum(1 for view in views if view.get("completed_viewing")) / total_views) * 100
    download_rate = (sum(1 for view in views if view.get("downloaded")) / total_views) * 100
    sign_rate = (sum(1 for view in views if view.get("signed")) / total_views) * 100
    
    return {
        "views": total_views,
        "unique_viewers": unique_viewers,
        "avg_time": avg_time,
        "completion_rate": completion_rate,
        "download_rate": download_rate,
        "sign_rate": sign_rate,
        "engagement": {
            "total_page_views": sum(len(view.get("pages_viewed", [])) for view in views),
            "avg_pages_per_session": sum(len(view.get("pages_viewed", [])) for view in views) / total_views if total_views > 0 else 0
        }
    }

# ==================== BASIC ENDPOINTS ====================

@api_router.get("/")
async def root():
    return {"message": "Document Platform API", "version": "1.0.0"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.utcnow()}

# Include the router in the main app
app.include_router(api_router)

# ==================== STARTUP/SHUTDOWN EVENTS ====================

@app.on_event("startup")
async def startup_event():
    """Initialize database and create indexes"""
    logger.info("Starting up Document Platform API")
    await create_indexes()
    logger.info("Database indexes created")

@app.on_event("shutdown")
async def shutdown_event():
    """Cleanup database connections"""
    logger.info("Shutting down Document Platform API")
    await Database.close_connection()
