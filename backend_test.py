import requests
import json
import unittest
import uuid
import os
from datetime import datetime

# Get the backend URL from the frontend .env file
BACKEND_URL = "https://7056c40d-c306-42f0-bfc6-379013315fc8.preview.emergentagent.com/api"

def test_backend_api():
    """Test the backend API endpoints"""
    print("Starting API tests...")
    
    # Test variables
    base_url = BACKEND_URL
    headers = {"Content-Type": "application/json"}
    test_user_email = f"test.user.{uuid.uuid4()}@example.com"
    test_user_password = "SecurePassword123!"
    test_document_id = None
    test_section_id = None
    
    # 1. Test health check endpoint
    print("\n1. Testing health check endpoint...")
    response = requests.get(f"{base_url}/health")
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    assert data["status"] == "healthy", f"Expected 'healthy', got {data['status']}"
    print("✅ Health check endpoint working")
    
    # 2. Test root endpoint
    print("\n2. Testing root endpoint...")
    response = requests.get(f"{base_url}/")
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    assert data["message"] == "Document Platform API", f"Expected 'Document Platform API', got {data['message']}"
    print("✅ Root endpoint working")
    
    # 3. Test user registration
    print("\n3. Testing user registration...")
    user_data = {
        "email": test_user_email,
        "full_name": "Test User",
        "role": "editor",
        "organization": "Test Organization",
        "password": test_user_password
    }
    
    response = requests.post(f"{base_url}/auth/register", json=user_data)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify response structure
    assert "access_token" in data, "Missing access_token in response"
    assert data["token_type"] == "bearer", f"Expected 'bearer', got {data['token_type']}"
    assert "user" in data, "Missing user in response"
    assert data["user"]["email"] == test_user_email, f"Expected {test_user_email}, got {data['user']['email']}"
    
    # Save token for subsequent tests
    access_token = data["access_token"]
    headers["Authorization"] = f"Bearer {access_token}"
    print("✅ User registration working")
    
    # 4. Test user login
    print("\n4. Testing user login...")
    response = requests.post(
        f"{base_url}/auth/login",
        params={"email": test_user_email, "password": test_user_password}
    )
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify response structure
    assert "access_token" in data, "Missing access_token in response"
    assert data["token_type"] == "bearer", f"Expected 'bearer', got {data['token_type']}"
    assert "user" in data, "Missing user in response"
    assert data["user"]["email"] == test_user_email, f"Expected {test_user_email}, got {data['user']['email']}"
    
    # Update token for subsequent tests
    access_token = data["access_token"]
    headers["Authorization"] = f"Bearer {access_token}"
    print("✅ User login working")
    
    # 5. Test getting current user profile
    print("\n5. Testing get current user profile...")
    response = requests.get(f"{base_url}/users/me", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify user data
    assert data["email"] == test_user_email, f"Expected {test_user_email}, got {data['email']}"
    assert data["full_name"] == "Test User", f"Expected 'Test User', got {data['full_name']}"
    
    # Verify email integration fields
    assert "email_connections" in data, "Missing email_connections in user profile"
    assert "notification_settings" in data, "Missing notification_settings in user profile"
    assert "email_signature" in data, "Missing email_signature in user profile"
    print("✅ Get current user profile working with email fields")
    
    # 6. Test adding email connection
    print("\n6. Testing add email connection...")
    email_connection_data = {
        "provider": "gmail",
        "email_address": "test.gmail@example.com",
        "display_name": "Test Gmail",
        "is_primary": True
    }
    
    response = requests.post(f"{base_url}/users/me/email-connections", json=email_connection_data, headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify email connection data
    assert data["provider"] == "gmail", f"Expected 'gmail', got {data['provider']}"
    assert data["email_address"] == "test.gmail@example.com", f"Expected 'test.gmail@example.com', got {data['email_address']}"
    assert data["is_primary"] == True, f"Expected True, got {data['is_primary']}"
    
    # Save connection ID for subsequent tests
    gmail_connection_id = data["id"]
    print("✅ Add email connection working")
    
    # 7. Test adding second email connection
    print("\n7. Testing add second email connection...")
    email_connection_data = {
        "provider": "outlook",
        "email_address": "test.outlook@example.com",
        "display_name": "Test Outlook",
        "is_primary": False
    }
    
    response = requests.post(f"{base_url}/users/me/email-connections", json=email_connection_data, headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify email connection data
    assert data["provider"] == "outlook", f"Expected 'outlook', got {data['provider']}"
    assert data["email_address"] == "test.outlook@example.com", f"Expected 'test.outlook@example.com', got {data['email_address']}"
    assert data["is_primary"] == False, f"Expected False, got {data['is_primary']}"
    
    # Save connection ID for subsequent tests
    outlook_connection_id = data["id"]
    print("✅ Add second email connection working")
    
    # 8. Test getting email connections
    print("\n8. Testing get email connections...")
    response = requests.get(f"{base_url}/users/me/email-connections", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify email connections list
    assert isinstance(data, list), f"Expected list, got {type(data)}"
    assert len(data) == 2, f"Expected 2 email connections, got {len(data)}"
    
    # Check if our test connections are in the list
    gmail_found = False
    outlook_found = False
    for connection in data:
        if connection["provider"] == "gmail":
            gmail_found = True
            assert connection["is_primary"] == True, "Gmail connection should be primary"
        elif connection["provider"] == "outlook":
            outlook_found = True
            assert connection["is_primary"] == False, "Outlook connection should not be primary"
    
    assert gmail_found, "Gmail connection not found in connections list"
    assert outlook_found, "Outlook connection not found in connections list"
    print("✅ Get email connections working")
    
    # 9. Test setting primary email connection
    print("\n9. Testing set primary email connection...")
    response = requests.put(f"{base_url}/users/me/email-connections/{outlook_connection_id}/primary", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify response
    assert data["message"] == "Primary email connection updated", f"Expected 'Primary email connection updated', got {data['message']}"
    
    # Verify the change by getting connections again
    response = requests.get(f"{base_url}/users/me/email-connections", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    for connection in data:
        if connection["id"] == outlook_connection_id:
            assert connection["is_primary"] == True, "Outlook connection should now be primary"
        elif connection["id"] == gmail_connection_id:
            assert connection["is_primary"] == False, "Gmail connection should no longer be primary"
    
    print("✅ Set primary email connection working")
    
    # 10. Test updating notification settings
    print("\n10. Testing update notification settings...")
    notification_settings = {
        "email_notifications": True,
        "document_shared": True,
        "document_viewed": False,
        "comment_added": True,
        "mention_in_comment": True
    }
    
    response = requests.put(f"{base_url}/users/me/notification-settings", json=notification_settings, headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify response
    assert data["message"] == "Notification settings updated successfully", f"Expected 'Notification settings updated successfully', got {data['message']}"
    
    # Verify the change by getting user profile
    response = requests.get(f"{base_url}/users/me", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    assert data["notification_settings"]["document_viewed"] == False, "document_viewed setting should be False"
    assert data["notification_settings"]["comment_added"] == True, "comment_added setting should be True"
    print("✅ Update notification settings working")
    
    # 11. Test updating email signature
    print("\n11. Testing update email signature...")
    signature_data = {
        "signature": "Best regards,\nTest User\nTechCorp"
    }
    
    response = requests.put(f"{base_url}/users/me/email-signature", json=signature_data, headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify response
    assert data["message"] == "Email signature updated successfully", f"Expected 'Email signature updated successfully', got {data['message']}"
    
    # Verify the change by getting user profile
    response = requests.get(f"{base_url}/users/me", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    assert data["email_signature"] == "Best regards,\nTest User\nTechCorp", f"Email signature not updated correctly"
    print("✅ Update email signature working")
    
    # 12. Test removing email connection
    print("\n12. Testing remove email connection...")
    response = requests.delete(f"{base_url}/users/me/email-connections/{gmail_connection_id}", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify response
    assert data["message"] == "Email connection removed successfully", f"Expected 'Email connection removed successfully', got {data['message']}"
    
    # Verify the change by getting connections again
    response = requests.get(f"{base_url}/users/me/email-connections", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    assert len(data) == 1, f"Expected 1 email connection, got {len(data)}"
    assert data[0]["id"] == outlook_connection_id, "Wrong email connection was removed"
    print("✅ Remove email connection working")
    
    # 13. Test document creation
    print("\n13. Testing document creation...")
    document_data = {
        "title": "Test Document",
        "type": "rfp",
        "organization": "Test Organization",
        "sections": [
            {
                "title": "Introduction",
                "content": "This is a test document introduction.",
                "order": 1
            },
            {
                "title": "Requirements",
                "content": "These are the test requirements.",
                "order": 2
            }
        ],
        "collaborators": [],
        "tags": ["test", "api"],
        "metadata": {"purpose": "testing"}
    }
    
    response = requests.post(f"{base_url}/documents", json=document_data, headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify document data
    assert data["title"] == "Test Document", f"Expected 'Test Document', got {data['title']}"
    assert data["type"] == "rfp", f"Expected 'rfp', got {data['type']}"
    assert len(data["sections"]) == 2, f"Expected 2 sections, got {len(data['sections'])}"
    
    # Save document ID for subsequent tests
    test_document_id = data["id"]
    test_section_id = data["sections"][0]["id"]
    print("✅ Document creation working")
    
    # 14. Test sending document via email
    print("\n14. Testing send document via email...")
    email_data = {
        "recipients": ["recipient@example.com"],
        "subject": "Test Document",
        "message": "Please review this document."
    }
    
    response = requests.post(f"{base_url}/documents/{test_document_id}/send-email", json=email_data, headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify response
    assert data["message"] == "Email sent successfully", f"Expected 'Email sent successfully', got {data['message']}"
    assert data["recipients"] == ["recipient@example.com"], f"Expected ['recipient@example.com'], got {data['recipients']}"
    assert data["document_title"] == "Test Document", f"Expected 'Test Document', got {data['document_title']}"
    print("✅ Send document via email working")
    
    # 15. Test getting user documents
    print("\n15. Testing get documents...")
    response = requests.get(f"{base_url}/documents", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify documents list
    assert isinstance(data, list), f"Expected list, got {type(data)}"
    assert len(data) >= 1, f"Expected at least 1 document, got {len(data)}"
    
    # Check if our test document is in the list
    document_found = False
    for doc in data:
        if doc["id"] == test_document_id:
            document_found = True
            break
    
    assert document_found, "Test document not found in documents list"
    print("✅ Get documents working")
    
    # 16. Test getting a specific document
    print("\n16. Testing get specific document...")
    response = requests.get(f"{base_url}/documents/{test_document_id}", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify document data
    assert data["id"] == test_document_id, f"Expected {test_document_id}, got {data['id']}"
    assert data["title"] == "Test Document", f"Expected 'Test Document', got {data['title']}"
    print("✅ Get specific document working")
    
    # 17. Test updating a document
    print("\n17. Testing update document...")
    update_data = {
        "title": "Updated Test Document",
        "tags": ["test", "api", "updated"]
    }
    
    response = requests.put(f"{base_url}/documents/{test_document_id}", json=update_data, headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify updated document data
    assert data["id"] == test_document_id, f"Expected {test_document_id}, got {data['id']}"
    assert data["title"] == "Updated Test Document", f"Expected 'Updated Test Document', got {data['title']}"
    assert data["tags"] == ["test", "api", "updated"], f"Expected ['test', 'api', 'updated'], got {data['tags']}"
    print("✅ Update document working")
    
    # 18. Test adding a comment to a document
    print("\n18. Testing add comment...")
    comment_data = {
        "content": "This is a test comment."
    }
    
    response = requests.post(f"{base_url}/documents/{test_document_id}/comments", json=comment_data, headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify comment data
    assert data["content"] == "This is a test comment.", f"Expected 'This is a test comment.', got {data['content']}"
    assert "id" in data, "Missing id in response"
    print("✅ Add comment working")
    
    # 19. Test getting comments for a document
    print("\n19. Testing get comments...")
    response = requests.get(f"{base_url}/documents/{test_document_id}/comments", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify comments list
    assert isinstance(data, list), f"Expected list, got {type(data)}"
    assert len(data) >= 1, f"Expected at least 1 comment, got {len(data)}"
    
    # Check if our test comment is in the list
    comment_found = False
    for comment in data:
        if comment["content"] == "This is a test comment.":
            comment_found = True
            break
    
    assert comment_found, "Test comment not found in comments list"
    print("✅ Get comments working")
    
    # 20. Test tracking a document view
    print("\n20. Testing track document view...")
    view_data = {
        "document_id": test_document_id,
        "viewer_info": {
            "ip_address": "192.168.1.1",
            "user_agent": "Test User Agent"
        },
        "pages_viewed": [],
        "total_time_spent": 60,
        "completed_viewing": True
    }
    
    response = requests.post(f"{base_url}/tracking/view", json=view_data)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify response
    assert data["message"] == "View tracked successfully", f"Expected 'View tracked successfully', got {data['message']}"
    print("✅ Track document view working")
    
    # 21. Test getting document analytics
    print("\n21. Testing get document analytics...")
    response = requests.get(f"{base_url}/documents/{test_document_id}/analytics", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify analytics data structure
    assert "views" in data, "Missing views in response"
    assert "unique_viewers" in data, "Missing unique_viewers in response"
    assert "avg_time" in data, "Missing avg_time in response"
    print("✅ Get document analytics working")
    
    # 22. Test AI-powered RFP generation
    print("\n22. Testing AI RFP generation...")
    rfp_data = {
        "project_type": "Website Development",
        "industry": "technology",
        "budget_range": "50k-100k",
        "timeline": "6 months",
        "requirements": "Build a modern e-commerce website with user authentication, product catalog, shopping cart, and payment processing",
        "company_info": "TechCorp is a growing startup in the technology sector",
        "specific_deliverables": ["Frontend website", "Backend API", "Admin dashboard", "Mobile responsive design"],
        "evaluation_criteria": ["Technical expertise", "Portfolio quality", "Timeline adherence", "Cost effectiveness"],
        "additional_context": "This is a high-priority project for our company's growth strategy"
    }
    
    response = requests.post(f"{base_url}/ai/generate-rfp", json=rfp_data, headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify RFP generation response
    assert "sections" in data, "Missing sections in response"
    assert "message" in data, "Missing message in response"
    assert data["message"] == "RFP content generated successfully", f"Expected 'RFP content generated successfully', got {data['message']}"
    assert len(data["sections"]) >= 1, f"Expected at least 1 section, got {len(data['sections'])}"
    
    # Print section titles to verify content
    print("\nGenerated RFP Sections:")
    for section in data["sections"]:
        print(f"- {section['title']}")
    
    print("✅ AI RFP generation working")
    
    # 23. Test document performance analysis
    print("\n23. Testing document performance analysis...")
    response = requests.post(f"{base_url}/ai/analyze-document/{test_document_id}", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify analysis response
    assert "recommendations" in data, "Missing recommendations in response"
    assert "performance_data" in data, "Missing performance_data in response"
    print("✅ Document performance analysis working")
    
    # 24. Test document upload with text file
    print("\n24. Testing document upload with text file...")
    
    # Create a test text file
    test_file_path = "/tmp/test_document.txt"
    with open(test_file_path, "w") as f:
        f.write("This is a test proposal document. Section 1: Project Overview. This section describes the project details. Section 2: Requirements. This section outlines the requirements.")
    
    # Upload the file
    with open(test_file_path, "rb") as f:
        files = {"file": ("test_document.txt", f, "text/plain")}
        response = requests.post(
            f"{base_url}/documents/upload",
            files=files,
            data={"title": "Test Document Upload"},
            headers={"Authorization": headers["Authorization"]}
        )
    
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify upload response
    assert data["message"] == "Document uploaded and processed successfully", f"Expected 'Document uploaded and processed successfully', got {data['message']}"
    assert "document" in data, "Missing document in response"
    assert data["document"]["title"] == "Test Document Upload", f"Expected 'Test Document Upload', got {data['document']['title']}"
    
    # Save uploaded document ID and section ID for subsequent tests
    uploaded_document_id = data["document"]["id"]
    uploaded_section_id = data["document"]["sections"][0]["id"]
    print("✅ Document upload with text file working")
    
    # 25. Test updating a document section
    print("\n25. Testing update document section...")
    section_update_data = {
        "title": "Updated Section Title",
        "content": "This is updated section content with rich text formatting."
    }
    
    response = requests.put(
        f"{base_url}/documents/{uploaded_document_id}/sections/{uploaded_section_id}",
        json=section_update_data,
        headers=headers
    )
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify section update response
    assert data["message"] == "Section updated successfully", f"Expected 'Section updated successfully', got {data['message']}"
    
    # Verify the update by getting the document
    response = requests.get(f"{base_url}/documents/{uploaded_document_id}", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    section_found = False
    for section in data["sections"]:
        if section["id"] == uploaded_section_id:
            section_found = True
            assert section["title"] == "Updated Section Title", f"Expected 'Updated Section Title', got {section['title']}"
            assert section["content"] == "This is updated section content with rich text formatting.", f"Content not updated correctly"
            break
    
    assert section_found, "Updated section not found in document"
    print("✅ Update document section working")
    
    # 26. Test adding multimedia element to a section
    print("\n26. Testing add multimedia element to section...")
    multimedia_data = {
        "type": "video",
        "url": "https://example.com/video.mp4",
        "title": "Project Demo Video",
        "description": "Overview of the project"
    }
    
    response = requests.post(
        f"{base_url}/documents/{uploaded_document_id}/sections/{uploaded_section_id}/multimedia",
        json=multimedia_data,
        headers=headers
    )
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify multimedia element response
    assert data["message"] == "Multimedia element added successfully", f"Expected 'Multimedia element added successfully', got {data['message']}"
    assert "element" in data, "Missing element in response"
    assert data["element"]["type"] == "video", f"Expected 'video', got {data['element']['type']}"
    assert data["element"]["url"] == "https://example.com/video.mp4", f"URL not set correctly"
    
    # Save multimedia element ID for verification
    multimedia_element_id = data["element"]["id"]
    
    # Verify the multimedia element was added by getting the document
    response = requests.get(f"{base_url}/documents/{uploaded_document_id}", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    multimedia_found = False
    for section in data["sections"]:
        if section["id"] == uploaded_section_id:
            for element in section["multimedia_elements"]:
                if element["id"] == multimedia_element_id:
                    multimedia_found = True
                    assert element["type"] == "video", f"Expected 'video', got {element['type']}"
                    assert element["url"] == "https://example.com/video.mp4", f"URL not set correctly"
                    break
    
    assert multimedia_found, "Multimedia element not found in document section"
    print("✅ Add multimedia element to section working")
    
    # 27. Test adding interactive element to a section
    print("\n27. Testing add interactive element to section...")
    interactive_data = {
        "type": "signature_field",
        "label": "Client Signature",
        "required": True,
        "position": {"x": 0.5, "y": 0.8, "page": 1}
    }
    
    response = requests.post(
        f"{base_url}/documents/{uploaded_document_id}/sections/{uploaded_section_id}/interactive",
        json=interactive_data,
        headers=headers
    )
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    # Verify interactive element response
    assert data["message"] == "Interactive element added successfully", f"Expected 'Interactive element added successfully', got {data['message']}"
    assert "element" in data, "Missing element in response"
    assert data["element"]["type"] == "signature_field", f"Expected 'signature_field', got {data['element']['type']}"
    assert data["element"]["label"] == "Client Signature", f"Label not set correctly"
    
    # Save interactive element ID for verification
    interactive_element_id = data["element"]["id"]
    
    # Verify the interactive element was added by getting the document
    response = requests.get(f"{base_url}/documents/{uploaded_document_id}", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    data = response.json()
    
    interactive_found = False
    for section in data["sections"]:
        if section["id"] == uploaded_section_id:
            for element in section["interactive_elements"]:
                if element["id"] == interactive_element_id:
                    interactive_found = True
                    assert element["type"] == "signature_field", f"Expected 'signature_field', got {element['type']}"
                    assert element["label"] == "Client Signature", f"Label not set correctly"
                    break
    
    assert interactive_found, "Interactive element not found in document section"
    print("✅ Add interactive element to section working")
    
    # 28. Test document deletion (cleanup)
    print("\n28. Testing document deletion (cleanup)...")
    
    # Delete the first test document
    response = requests.delete(f"{base_url}/documents/{test_document_id}", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    
    # Delete the uploaded test document
    response = requests.delete(f"{base_url}/documents/{uploaded_document_id}", headers=headers)
    assert response.status_code == 200, f"Expected 200, got {response.status_code}"
    
    # Clean up the test file
    if os.path.exists(test_file_path):
        os.remove(test_file_path)
    
    print("✅ Document deletion (cleanup) working")
    
    print("\n✅ All tests passed successfully!")

def test_enhanced_document_upload():
    """Test the enhanced document upload functionality with formatting preservation"""
    print("Starting enhanced document upload tests...")
    
    # Test variables
    base_url = BACKEND_URL
    headers = {"Content-Type": "application/json"}
    test_user_email = f"test.user.{uuid.uuid4()}@example.com"
    test_user_password = "SecurePassword123!"
    
    # 1. Register a test user
    print("\n1. Registering test user...")
    user_data = {
        "email": test_user_email,
        "full_name": "Test User",
        "role": "editor",
        "organization": "Test Organization",
        "password": test_user_password
    }
    
    response = requests.post(f"{base_url}/auth/register", json=user_data)
    assert response.status_code == 200, f"Failed to register test user: {response.text}"
    data = response.json()
    
    # Save token for subsequent tests
    access_token = data["access_token"]
    headers["Authorization"] = f"Bearer {access_token}"
    print("✅ User registration successful")
    
    # 2. Create test files
    print("\n2. Creating test files...")
    
    # Create a DOCX file with formatting
    try:
        from docx import Document
        import time
        
        docx_path = "/tmp/test_formatted_document.docx"
        doc = Document()
        
        # Add a title with formatting
        title = doc.add_heading('Formatted Test Document', 0)
        
        # Add a paragraph with bold and italic text
        p = doc.add_paragraph('This document contains ')
        p.add_run('bold text, ').bold = True
        p.add_run('italic text, ').italic = True
        p.add_run('and ')
        p.add_run('bold-italic text.').bold = True
        p.add_run('bold-italic text.').italic = True
        
        # Add a bulleted list
        doc.add_paragraph('List item 1', style='List Bullet')
        doc.add_paragraph('List item 2', style='List Bullet')
        doc.add_paragraph('List item 3', style='List Bullet')
        
        # Add a numbered list
        doc.add_paragraph('Numbered item 1', style='List Number')
        doc.add_paragraph('Numbered item 2', style='List Number')
        doc.add_paragraph('Numbered item 3', style='List Number')
        
        # Add a table
        table = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                table.cell(i, j).text = f"Cell {i+1},{j+1}"
        
        # Save the document
        doc.save(docx_path)
        
        # Create a plain text file for comparison
        txt_path = "/tmp/test_plain_document.txt"
        with open(txt_path, "w") as f:
            f.write("This is a plain text document.\n\nIt has multiple paragraphs but no formatting.\n\n- Item 1\n- Item 2\n- Item 3")
        
        print("✅ Test files created successfully")
        
        # 3. Test DOCX upload with formatting preservation
        print("\n3. Testing DOCX upload with formatting preservation...")
        
        # Upload the DOCX file
        with open(docx_path, "rb") as f:
            files = {"file": ("test_formatted_document.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                headers={"Authorization": headers["Authorization"]}
            )
        
        assert response.status_code == 200, f"Failed to upload DOCX file: {response.text}"
        data = response.json()
        
        # Verify response structure
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        
        # Verify HTML content generation
        assert "pages" in data["document"], "Missing pages in document"
        assert len(data["document"]["pages"]) > 0, "No pages found in document"
        
        # Check that content is HTML, not plain text
        page_content = data["document"]["pages"][0]["content"]
        assert "<div" in page_content, "Content is not in HTML format"
        
        # Check for formatting elements
        formatting_found = False
        for tag in ["<b>", "<strong>", "<i>", "<em>", "<ul>", "<ol>", "<li>", "<table>"]:
            if tag in page_content.lower():
                formatting_found = True
                break
        
        assert formatting_found, "HTML content does not contain formatting tags"
        
        # Save document ID for retrieval test
        docx_document_id = data["document"]["id"]
        
        # Print some details about the document
        print(f"Document ID: {docx_document_id}")
        print(f"Document title: {data['document']['title']}")
        print(f"Total pages: {data['document']['total_pages']}")
        print(f"Contains formatting: {data['document']['metadata']['contains_formatting']}")
        
        print("✅ DOCX upload with formatting preservation working")
        
        # 4. Test plain text upload with HTML conversion
        print("\n4. Testing plain text upload with HTML conversion...")
        
        # Upload the text file
        with open(txt_path, "rb") as f:
            files = {"file": ("test_plain_document.txt", f, "text/plain")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                headers={"Authorization": headers["Authorization"]}
            )
        
        assert response.status_code == 200, f"Failed to upload text file: {response.text}"
        data = response.json()
        
        # Verify response structure
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        
        # Verify HTML content generation for plain text
        assert "pages" in data["document"], "Missing pages in document"
        assert len(data["document"]["pages"]) > 0, "No pages found in document"
        
        # Check that content is HTML, not plain text
        page_content = data["document"]["pages"][0]["content"]
        assert "<div" in page_content, "Content is not in HTML format"
        
        # Save document ID for retrieval test
        txt_document_id = data["document"]["id"]
        
        # Print some details about the document
        print(f"Document ID: {txt_document_id}")
        print(f"Document title: {data['document']['title']}")
        print(f"Total pages: {data['document']['total_pages']}")
        
        print("✅ Plain text upload with HTML conversion working")
        
        # 5. Test document retrieval with HTML content
        print("\n5. Testing document retrieval with HTML content...")
        
        # Give the server a moment to process the document
        time.sleep(1)
        
        # Retrieve the DOCX document
        response = requests.get(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve DOCX document: {response.text}"
        data = response.json()
        
        # Verify document structure
        assert data["id"] == docx_document_id, "Document ID mismatch"
        assert "pages" in data, "Missing pages in retrieved document"
        assert len(data["pages"]) > 0, "No pages found in retrieved document"
        
        # Check that content is HTML
        page_content = data["pages"][0]["content"]
        assert "<div" in page_content, "Retrieved content is not in HTML format"
        
        # Check for formatting elements
        formatting_found = False
        for tag in ["<b>", "<strong>", "<i>", "<em>", "<ul>", "<ol>", "<li>", "<table>"]:
            if tag in page_content.lower():
                formatting_found = True
                break
        
        assert formatting_found, "Retrieved HTML content does not contain formatting tags"
        
        # Check metadata
        assert "metadata" in data, "Missing metadata in document"
        assert "contains_formatting" in data["metadata"], "Missing contains_formatting flag in metadata"
        assert data["metadata"]["contains_formatting"], "contains_formatting flag should be True"
        
        # Print some details about the retrieved document
        print(f"Retrieved document title: {data['title']}")
        print(f"Retrieved document total pages: {data['total_pages']}")
        print(f"Retrieved document contains formatting: {data['metadata']['contains_formatting']}")
        
        print("✅ Document retrieval with HTML content working")
        
        # 6. Test backward compatibility with plain text documents
        print("\n6. Testing backward compatibility with plain text documents...")
        
        # Retrieve the plain text document
        response = requests.get(f"{base_url}/documents/{txt_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve text document: {response.text}"
        data = response.json()
        
        # Verify document structure
        assert data["id"] == txt_document_id, "Document ID mismatch"
        
        # Check that both sections and pages are present
        assert "sections" in data, "Missing sections in retrieved document"
        assert "pages" in data, "Missing pages in retrieved document"
        
        # Verify content in both sections and pages
        assert len(data["sections"]) > 0, "No sections found in document"
        assert len(data["pages"]) > 0, "No pages found in document"
        
        # Check that content is HTML in both
        section_content = data["sections"][0]["content"]
        page_content = data["pages"][0]["content"]
        
        assert "<div" in section_content, "Section content is not in HTML format"
        assert "<div" in page_content, "Page content is not in HTML format"
        
        # Print some details about the retrieved document
        print(f"Retrieved document title: {data['title']}")
        print(f"Retrieved document has {len(data['sections'])} sections and {len(data['pages'])} pages")
        
        print("✅ Backward compatibility with plain text documents working")
        
        # 7. Test updating a document page
        print("\n7. Testing document page update...")
        
        # Update the first page of the DOCX document
        page_update_data = {
            "title": "Updated Page Title",
            "content": "<div style='font-family: Arial; color: blue;'>This is updated content with <b>HTML formatting</b> and <i>styling</i>.</div>"
        }
        
        response = requests.put(
            f"{base_url}/documents/{docx_document_id}/pages/1",
            json=page_update_data,
            headers=headers
        )
        
        assert response.status_code == 200, f"Failed to update document page: {response.text}"
        data = response.json()
        
        # Verify update response
        assert data["message"] == "Page updated successfully", "Unexpected response message"
        
        # Retrieve the document to verify the update
        response = requests.get(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve updated document: {response.text}"
        data = response.json()
        
        # Check that the page was updated
        assert data["pages"][0]["title"] == "Updated Page Title", "Page title not updated"
        assert "<b>HTML formatting</b>" in data["pages"][0]["content"], "HTML formatting not preserved in update"
        
        print("✅ Document page update with HTML formatting working")
        
        # 8. Test adding multimedia to a page
        print("\n8. Testing adding multimedia to a page...")
        
        # Add a multimedia element to the first page
        multimedia_data = {
            "type": "image",
            "url": "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==",
            "title": "Test Image",
            "description": "A test base64 image"
        }
        
        response = requests.post(
            f"{base_url}/documents/{docx_document_id}/pages/1/multimedia",
            json=multimedia_data,
            headers=headers
        )
        
        assert response.status_code == 200, f"Failed to add multimedia to page: {response.text}"
        data = response.json()
        
        # Verify multimedia response
        assert data["message"] == "Multimedia element added to page successfully", "Unexpected response message"
        assert "element" in data, "Missing element in response"
        assert data["element"]["type"] == "image", "Element type mismatch"
        
        # Retrieve the document to verify the multimedia element
        response = requests.get(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve document with multimedia: {response.text}"
        data = response.json()
        
        # Check that the multimedia element was added
        assert len(data["pages"][0]["multimedia_elements"]) > 0, "No multimedia elements found"
        assert data["pages"][0]["multimedia_elements"][0]["type"] == "image", "Multimedia element type mismatch"
        
        print("✅ Adding multimedia to page working")
        
        # 9. Clean up
        print("\n9. Cleaning up test documents...")
        
        # Delete the test documents
        response = requests.delete(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to delete DOCX document: {response.text}"
        
        response = requests.delete(f"{base_url}/documents/{txt_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to delete text document: {response.text}"
        
        # Clean up the test files
        if os.path.exists(docx_path):
            os.remove(docx_path)
        
        if os.path.exists(txt_path):
            os.remove(txt_path)
        
        print("✅ Test cleanup successful")
        
        print("\n✅ All enhanced document upload tests passed successfully!")
    
    except ImportError:
        print("⚠️ python-docx library not available. Skipping enhanced document upload tests.")
        print("To run these tests, install the required libraries: pip install python-docx Pillow")

if __name__ == "__main__":
    # Run the enhanced document upload tests
    test_enhanced_document_upload()