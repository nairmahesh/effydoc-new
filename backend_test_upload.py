import requests
import json
import unittest
import uuid
import os
import time
from datetime import datetime

# Get the backend URL from the frontend .env file
BACKEND_URL = "https://7056c40d-c306-42f0-bfc6-379013315fc8.preview.emergentagent.com/api"

def test_document_upload_functionality():
    """Test the document upload functionality with various file types and extract_text parameter"""
    print("Starting document upload tests...")
    
    # Test variables
    base_url = BACKEND_URL
    headers = {"Content-Type": "application/json"}
    test_user_email = f"test.user.{uuid.uuid4()}@example.com"
    test_user_password = "SecurePassword123!"
    
    # 1. Register a test user
    print("\n1. Registering test user...")
    user_data = {
        "email": test_user_email,
        "full_name": "Test User",
        "role": "editor",
        "organization": "Test Organization",
        "password": test_user_password
    }
    
    response = requests.post(f"{base_url}/auth/register", json=user_data)
    assert response.status_code == 200, f"Failed to register test user: {response.text}"
    data = response.json()
    
    # Save token for subsequent tests
    access_token = data["access_token"]
    auth_headers = {"Authorization": f"Bearer {access_token}"}
    print("✅ User registration successful")
    
    # 2. Create test files
    print("\n2. Creating test files...")
    
    # Create a text file
    txt_path = "/tmp/test_document.txt"
    with open(txt_path, "w") as f:
        f.write("This is a test document.\n\nIt has multiple paragraphs.\n\n- Item 1\n- Item 2\n- Item 3")
    
    # Create a PDF file (simple text-based PDF)
    pdf_path = None
    try:
        from reportlab.pdfgen import canvas
        pdf_path = "/tmp/test_document.pdf"
        c = canvas.Canvas(pdf_path)
        c.drawString(100, 750, "Test PDF Document")
        c.drawString(100, 730, "This is a test PDF file for upload testing.")
        c.drawString(100, 710, "It contains multiple lines of text.")
        c.drawString(100, 690, "- Item 1")
        c.drawString(100, 670, "- Item 2")
        c.drawString(100, 650, "- Item 3")
        c.save()
        print("Created PDF test file")
    except ImportError:
        print("reportlab not installed, skipping PDF test file creation")
    
    # Create a DOCX file if python-docx is available
    docx_path = None
    try:
        from docx import Document
        docx_path = "/tmp/test_document.docx"
        doc = Document()
        doc.add_heading('Test DOCX Document', 0)
        p = doc.add_paragraph('This is a test DOCX file for upload testing. ')
        p.add_run('It contains formatted text.').bold = True
        doc.add_paragraph('It has multiple paragraphs.').italic = True
        doc.add_paragraph('- Item 1', style='List Bullet')
        doc.add_paragraph('- Item 2', style='List Bullet')
        doc.add_paragraph('- Item 3', style='List Bullet')
        doc.save(docx_path)
        print("Created DOCX test file")
    except ImportError:
        print("python-docx not installed, skipping DOCX test file creation")
    
    print("✅ Test files created")
    
    # 3. Test text file upload with extract_text=True (default)
    print("\n3. Testing text file upload with extract_text=True (default)...")
    
    with open(txt_path, "rb") as f:
        files = {"file": ("test_document.txt", f, "text/plain")}
        response = requests.post(
            f"{base_url}/documents/upload",
            files=files,
            headers=auth_headers
        )
    
    assert response.status_code == 200, f"Failed to upload text file with extract_text=True: {response.text}"
    data = response.json()
    
    # Verify response
    assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
    assert "document" in data, "Missing document in response"
    assert "pages" in data["document"], "Missing pages in document"
    assert len(data["document"]["pages"]) > 0, "No pages found in document"
    
    # Check that content is HTML
    page_content = data["document"]["pages"][0]["content"]
    assert "<div" in page_content, "Content is not in HTML format"
    
    # Save document ID
    txt_document_id = data["document"]["id"]
    print(f"Text file uploaded successfully with ID: {txt_document_id}")
    print("✅ Text file upload with extract_text=True working")
    
    # 4. Test PDF file upload with extract_text=True (default)
    if pdf_path:
        print("\n4. Testing PDF file upload with extract_text=True (default)...")
        
        with open(pdf_path, "rb") as f:
            files = {"file": ("test_document.pdf", f, "application/pdf")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                headers=auth_headers
            )
        
        assert response.status_code == 200, f"Failed to upload PDF file with extract_text=True: {response.text}"
        data = response.json()
        
        # Verify response
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        assert "pages" in data["document"], "Missing pages in document"
        assert len(data["document"]["pages"]) > 0, "No pages found in document"
        
        # Check that content is HTML
        page_content = data["document"]["pages"][0]["content"]
        assert "<div" in page_content, "Content is not in HTML format"
        
        # Save document ID
        pdf_extract_document_id = data["document"]["id"]
        print(f"PDF file uploaded with extract_text=True successfully with ID: {pdf_extract_document_id}")
        print("✅ PDF file upload with extract_text=True working")
    else:
        print("⚠️ Skipping PDF file upload with extract_text=True test (no PDF file created)")
        pdf_extract_document_id = None
    
    # 5. Test PDF file upload with extract_text=False
    if pdf_path:
        print("\n5. Testing PDF file upload with extract_text=False...")
        
        with open(pdf_path, "rb") as f:
            files = {"file": ("test_document.pdf", f, "application/pdf")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                data={"extract_text": "false"},
                headers=auth_headers
            )
        
        assert response.status_code == 200, f"Failed to upload PDF file with extract_text=False: {response.text}"
        data = response.json()
        
        # Verify response
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        assert "pages" in data["document"], "Missing pages in document"
        assert len(data["document"]["pages"]) > 0, "No pages found in document"
        
        # Check that content contains PDF embed
        page_content = data["document"]["pages"][0]["content"]
        assert "<embed" in page_content, "Content does not contain PDF embed"
        assert "data:application/pdf;base64," in page_content, "Content does not contain base64 PDF data"
        
        # Save document ID
        pdf_preserve_document_id = data["document"]["id"]
        print(f"PDF file uploaded with extract_text=False successfully with ID: {pdf_preserve_document_id}")
        print("✅ PDF file upload with extract_text=False working")
    else:
        print("⚠️ Skipping PDF file upload with extract_text=False test (no PDF file created)")
        pdf_preserve_document_id = None
    
    # 6. Test DOCX file upload
    if docx_path:
        print("\n6. Testing DOCX file upload...")
        
        with open(docx_path, "rb") as f:
            files = {"file": ("test_document.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                headers=auth_headers
            )
        
        assert response.status_code == 200, f"Failed to upload DOCX file: {response.text}"
        data = response.json()
        
        # Verify response
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        assert "pages" in data["document"], "Missing pages in document"
        assert len(data["document"]["pages"]) > 0, "No pages found in document"
        
        # Check that content is HTML with formatting
        page_content = data["document"]["pages"][0]["content"]
        assert "<div" in page_content, "Content is not in HTML format"
        
        # Check for formatting elements (may vary based on mammoth.js conversion)
        formatting_found = False
        for tag in ["<b>", "<strong>", "<i>", "<em>", "<ul>", "<ol>", "<li>", "<p>"]:
            if tag in page_content.lower():
                formatting_found = True
                break
        
        assert formatting_found, "HTML content does not contain formatting tags"
        
        # Save document ID
        docx_document_id = data["document"]["id"]
        print(f"DOCX file uploaded successfully with ID: {docx_document_id}")
        print("✅ DOCX file upload working")
    else:
        print("⚠️ Skipping DOCX file upload test (no DOCX file created)")
        docx_document_id = None
    
    # 7. Test retrieving uploaded documents
    print("\n7. Testing document retrieval...")
    
    # Retrieve text document
    response = requests.get(f"{base_url}/documents/{txt_document_id}", headers=auth_headers)
    assert response.status_code == 200, f"Failed to retrieve text document: {response.text}"
    data = response.json()
    
    # Verify document structure
    assert data["id"] == txt_document_id, "Document ID mismatch"
    assert "pages" in data, "Missing pages in retrieved document"
    assert len(data["pages"]) > 0, "No pages found in retrieved document"
    
    # Check that content is HTML
    page_content = data["pages"][0]["content"]
    assert "<div" in page_content, "Retrieved content is not in HTML format"
    
    print("✅ Text document retrieval working")
    
    # Retrieve PDF document with extract_text=True
    if pdf_extract_document_id:
        response = requests.get(f"{base_url}/documents/{pdf_extract_document_id}", headers=auth_headers)
        assert response.status_code == 200, f"Failed to retrieve PDF document with extract_text=True: {response.text}"
        data = response.json()
        
        # Verify document structure
        assert data["id"] == pdf_extract_document_id, "Document ID mismatch"
        assert "pages" in data, "Missing pages in retrieved document"
        assert len(data["pages"]) > 0, "No pages found in retrieved document"
        
        # Check that content is HTML
        page_content = data["pages"][0]["content"]
        assert "<div" in page_content, "Retrieved content is not in HTML format"
        
        print("✅ PDF document with extract_text=True retrieval working")
    
    # Retrieve PDF document with extract_text=False
    if pdf_preserve_document_id:
        response = requests.get(f"{base_url}/documents/{pdf_preserve_document_id}", headers=auth_headers)
        assert response.status_code == 200, f"Failed to retrieve PDF document with extract_text=False: {response.text}"
        data = response.json()
        
        # Verify document structure
        assert data["id"] == pdf_preserve_document_id, "Document ID mismatch"
        assert "pages" in data, "Missing pages in retrieved document"
        assert len(data["pages"]) > 0, "No pages found in retrieved document"
        
        # Check that content contains PDF embed
        page_content = data["pages"][0]["content"]
        assert "<embed" in page_content, "Retrieved content does not contain PDF embed"
        assert "data:application/pdf;base64," in page_content, "Retrieved content does not contain base64 PDF data"
        
        print("✅ PDF document with extract_text=False retrieval working")
    
    # Retrieve DOCX document
    if docx_document_id:
        response = requests.get(f"{base_url}/documents/{docx_document_id}", headers=auth_headers)
        assert response.status_code == 200, f"Failed to retrieve DOCX document: {response.text}"
        data = response.json()
        
        # Verify document structure
        assert data["id"] == docx_document_id, "Document ID mismatch"
        assert "pages" in data, "Missing pages in retrieved document"
        assert len(data["pages"]) > 0, "No pages found in retrieved document"
        
        # Check that content is HTML with formatting
        page_content = data["pages"][0]["content"]
        assert "<div" in page_content, "Retrieved content is not in HTML format"
        
        print("✅ DOCX document retrieval working")
    
    # 8. Test uploading an unsupported file type
    print("\n8. Testing upload of unsupported file type...")
    
    # Create a test file with unsupported extension
    unsupported_path = "/tmp/test_document.xyz"
    with open(unsupported_path, "w") as f:
        f.write("This is an unsupported file type.")
    
    with open(unsupported_path, "rb") as f:
        files = {"file": ("test_document.xyz", f, "application/octet-stream")}
        response = requests.post(
            f"{base_url}/documents/upload",
            files=files,
            headers=auth_headers
        )
    
    assert response.status_code == 400, f"Expected 400 status code for unsupported file type, got {response.status_code}"
    data = response.json()
    assert "detail" in data, "Missing error detail in response"
    assert "Unsupported file type" in data["detail"], "Unexpected error message"
    
    print("✅ Unsupported file type rejection working")
    
    # 9. Test uploading with invalid extract_text parameter
    print("\n9. Testing upload with invalid extract_text parameter...")
    
    with open(txt_path, "rb") as f:
        files = {"file": ("test_document.txt", f, "text/plain")}
        response = requests.post(
            f"{base_url}/documents/upload",
            files=files,
            data={"extract_text": "invalid_value"},
            headers=auth_headers
        )
    
    # This should still work as FastAPI will convert "invalid_value" to True
    assert response.status_code == 200, f"Failed to upload with invalid extract_text parameter: {response.text}"
    
    print("✅ Upload with invalid extract_text parameter handling working")
    
    # 10. Test uploading without authentication
    print("\n10. Testing upload without authentication...")
    
    with open(txt_path, "rb") as f:
        files = {"file": ("test_document.txt", f, "text/plain")}
        response = requests.post(
            f"{base_url}/documents/upload",
            files=files
        )
    
    assert response.status_code == 401 or response.status_code == 403, f"Expected 401 or 403 status code for unauthenticated upload, got {response.status_code}"
    
    print("✅ Unauthenticated upload rejection working")
    
    # 11. Clean up
    print("\n11. Cleaning up test documents...")
    
    # Delete the test documents
    response = requests.delete(f"{base_url}/documents/{txt_document_id}", headers=auth_headers)
    assert response.status_code == 200, f"Failed to delete text document: {response.text}"
    
    if pdf_extract_document_id:
        response = requests.delete(f"{base_url}/documents/{pdf_extract_document_id}", headers=auth_headers)
        assert response.status_code == 200, f"Failed to delete PDF document with extract_text=True: {response.text}"
    
    if pdf_preserve_document_id:
        response = requests.delete(f"{base_url}/documents/{pdf_preserve_document_id}", headers=auth_headers)
        assert response.status_code == 200, f"Failed to delete PDF document with extract_text=False: {response.text}"
    
    if docx_document_id:
        response = requests.delete(f"{base_url}/documents/{docx_document_id}", headers=auth_headers)
        assert response.status_code == 200, f"Failed to delete DOCX document: {response.text}"
    
    # Clean up the test files
    if os.path.exists(txt_path):
        os.remove(txt_path)
    
    if pdf_path and os.path.exists(pdf_path):
        os.remove(pdf_path)
    
    if docx_path and os.path.exists(docx_path):
        os.remove(docx_path)
    
    if os.path.exists(unsupported_path):
        os.remove(unsupported_path)
    
    print("✅ Test cleanup successful")
    
    print("\n✅ All document upload tests passed successfully!")

if __name__ == "__main__":
    test_document_upload_functionality()