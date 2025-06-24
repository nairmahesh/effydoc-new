import requests
import json
import unittest
import uuid
import os
from datetime import datetime
import base64
from io import BytesIO
from docx import Document
import tempfile
import time

# Get the backend URL from the frontend .env file
BACKEND_URL = "https://4aa54517-8e64-4037-8338-35ac94bb76b8.preview.emergentagent.com/api"

class EnhancedUploadTest(unittest.TestCase):
    """Test the enhanced document upload functionality"""
    
    def setUp(self):
        """Set up test environment"""
        self.base_url = BACKEND_URL
        self.headers = {"Content-Type": "application/json"}
        
        # Create a test user for authentication
        self.test_user_email = f"test.user.{uuid.uuid4()}@example.com"
        self.test_user_password = "SecurePassword123!"
        
        # Register test user
        user_data = {
            "email": self.test_user_email,
            "full_name": "Test User",
            "role": "editor",
            "organization": "Test Organization",
            "password": self.test_user_password
        }
        
        response = requests.post(f"{self.base_url}/auth/register", json=user_data)
        self.assertEqual(response.status_code, 200, f"Failed to register test user: {response.text}")
        
        data = response.json()
        self.access_token = data["access_token"]
        self.headers["Authorization"] = f"Bearer {self.access_token}"
        
        # Create test files
        self.create_test_files()
    
    def create_test_files(self):
        """Create test files for upload testing"""
        # Create a DOCX file with formatting
        self.docx_path = "/tmp/test_formatted_document.docx"
        doc = Document()
        
        # Add a title with formatting
        title = doc.add_heading('Formatted Test Document', 0)
        
        # Add a paragraph with bold and italic text
        p = doc.add_paragraph('This document contains ')
        p.add_run('bold text, ').bold = True
        p.add_run('italic text, ').italic = True
        p.add_run('and ')
        p.add_run('bold-italic text.').bold = True
        p.add_run('bold-italic text.').italic = True
        
        # Add a bulleted list
        doc.add_paragraph('List item 1', style='List Bullet')
        doc.add_paragraph('List item 2', style='List Bullet')
        doc.add_paragraph('List item 3', style='List Bullet')
        
        # Add a numbered list
        doc.add_paragraph('Numbered item 1', style='List Number')
        doc.add_paragraph('Numbered item 2', style='List Number')
        doc.add_paragraph('Numbered item 3', style='List Number')
        
        # Add a table
        table = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                table.cell(i, j).text = f"Cell {i+1},{j+1}"
        
        # Save the document
        doc.save(self.docx_path)
        
        # Create a plain text file for comparison
        self.txt_path = "/tmp/test_plain_document.txt"
        with open(self.txt_path, "w") as f:
            f.write("This is a plain text document.\n\nIt has multiple paragraphs but no formatting.\n\n- Item 1\n- Item 2\n- Item 3")
    
    def test_docx_upload_with_formatting(self):
        """Test uploading a DOCX file with formatting preservation"""
        print("\nTesting DOCX upload with formatting preservation...")
        
        # Upload the DOCX file
        with open(self.docx_path, "rb") as f:
            files = {"file": ("test_formatted_document.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post(
                f"{self.base_url}/documents/upload",
                files=files,
                headers={"Authorization": self.headers["Authorization"]}
            )
        
        self.assertEqual(response.status_code, 200, f"Failed to upload DOCX file: {response.text}")
        data = response.json()
        
        # Verify response structure
        self.assertEqual(data["message"], "Document uploaded and processed successfully", 
                         "Unexpected response message")
        self.assertIn("document", data, "Missing document in response")
        
        # Verify HTML content generation
        self.assertIn("pages", data["document"], "Missing pages in document")
        self.assertTrue(len(data["document"]["pages"]) > 0, "No pages found in document")
        
        # Check that content is HTML, not plain text
        page_content = data["document"]["pages"][0]["content"]
        self.assertIn("<div", page_content, "Content is not in HTML format")
        
        # Check for formatting elements
        self.assertTrue(any(tag in page_content.lower() for tag in ["<b>", "<strong>", "<i>", "<em>", "<ul>", "<ol>", "<li>", "<table>"]), 
                        "HTML content does not contain formatting tags")
        
        # Save document ID for retrieval test
        self.docx_document_id = data["document"]["id"]
        
        print("✅ DOCX upload with formatting preservation working")
    
    def test_text_upload_with_html_conversion(self):
        """Test uploading a plain text file with HTML conversion"""
        print("\nTesting plain text upload with HTML conversion...")
        
        # Upload the text file
        with open(self.txt_path, "rb") as f:
            files = {"file": ("test_plain_document.txt", f, "text/plain")}
            response = requests.post(
                f"{self.base_url}/documents/upload",
                files=files,
                headers={"Authorization": self.headers["Authorization"]}
            )
        
        self.assertEqual(response.status_code, 200, f"Failed to upload text file: {response.text}")
        data = response.json()
        
        # Verify response structure
        self.assertEqual(data["message"], "Document uploaded and processed successfully", 
                         "Unexpected response message")
        self.assertIn("document", data, "Missing document in response")
        
        # Verify HTML content generation for plain text
        self.assertIn("pages", data["document"], "Missing pages in document")
        self.assertTrue(len(data["document"]["pages"]) > 0, "No pages found in document")
        
        # Check that content is HTML, not plain text
        page_content = data["document"]["pages"][0]["content"]
        self.assertIn("<div", page_content, "Content is not in HTML format")
        
        # Save document ID for retrieval test
        self.txt_document_id = data["document"]["id"]
        
        print("✅ Plain text upload with HTML conversion working")
    
    def test_document_retrieval_with_html_content(self):
        """Test retrieving documents with HTML content"""
        print("\nTesting document retrieval with HTML content...")
        
        # First run the upload test to get a document ID
        if not hasattr(self, 'docx_document_id'):
            self.test_docx_upload_with_formatting()
            
        # Give the server a moment to process the document
        time.sleep(1)
        
        # Retrieve the DOCX document
        response = requests.get(f"{self.base_url}/documents/{self.docx_document_id}", 
                               headers=self.headers)
        
        self.assertEqual(response.status_code, 200, f"Failed to retrieve DOCX document: {response.text}")
        data = response.json()
        
        # Verify document structure
        self.assertEqual(data["id"], self.docx_document_id, "Document ID mismatch")
        self.assertIn("pages", data, "Missing pages in retrieved document")
        self.assertTrue(len(data["pages"]) > 0, "No pages found in retrieved document")
        
        # Check that content is HTML
        page_content = data["pages"][0]["content"]
        self.assertIn("<div", page_content, "Retrieved content is not in HTML format")
        
        # Check for formatting elements
        self.assertTrue(any(tag in page_content.lower() for tag in ["<b>", "<strong>", "<i>", "<em>", "<ul>", "<ol>", "<li>", "<table>"]), 
                        "Retrieved HTML content does not contain formatting tags")
        
        # Check metadata
        self.assertIn("metadata", data, "Missing metadata in document")
        self.assertIn("contains_formatting", data["metadata"], "Missing contains_formatting flag in metadata")
        self.assertTrue(data["metadata"]["contains_formatting"], "contains_formatting flag should be True")
        
        print("✅ Document retrieval with HTML content working")
    
    def test_backward_compatibility(self):
        """Test backward compatibility with plain text documents"""
        print("\nTesting backward compatibility with plain text documents...")
        
        # Retrieve the plain text document
        response = requests.get(f"{self.base_url}/documents/{self.txt_document_id}", 
                               headers=self.headers)
        
        self.assertEqual(response.status_code, 200, f"Failed to retrieve text document: {response.text}")
        data = response.json()
        
        # Verify document structure
        self.assertEqual(data["id"], self.txt_document_id, "Document ID mismatch")
        
        # Check that both sections and pages are present
        self.assertIn("sections", data, "Missing sections in retrieved document")
        self.assertIn("pages", data, "Missing pages in retrieved document")
        
        # Verify content in both sections and pages
        self.assertTrue(len(data["sections"]) > 0, "No sections found in document")
        self.assertTrue(len(data["pages"]) > 0, "No pages found in document")
        
        # Check that content is HTML in both
        section_content = data["sections"][0]["content"]
        page_content = data["pages"][0]["content"]
        
        self.assertIn("<div", section_content, "Section content is not in HTML format")
        self.assertIn("<div", page_content, "Page content is not in HTML format")
        
        print("✅ Backward compatibility with plain text documents working")
    
    def tearDown(self):
        """Clean up after tests"""
        # Delete test documents
        if hasattr(self, 'docx_document_id'):
            requests.delete(f"{self.base_url}/documents/{self.docx_document_id}", 
                           headers=self.headers)
        
        if hasattr(self, 'txt_document_id'):
            requests.delete(f"{self.base_url}/documents/{self.txt_document_id}", 
                           headers=self.headers)
        
        # Delete test files
        if os.path.exists(self.docx_path):
            os.remove(self.docx_path)
        
        if os.path.exists(self.txt_path):
            os.remove(self.txt_path)

if __name__ == "__main__":
    unittest.main()