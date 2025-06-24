import requests
import json
import uuid
import os
from datetime import datetime
import time

# Get the backend URL from the frontend .env file
BACKEND_URL = "https://7056c40d-c306-42f0-bfc6-379013315fc8.preview.emergentagent.com/api"

def test_document_upload_and_viewing():
    """Test the document upload and viewing functionality with formatting preservation"""
    print("Starting document upload and viewing tests...")
    
    # Test variables
    base_url = BACKEND_URL
    headers = {"Content-Type": "application/json"}
    test_user_email = f"test.user.{uuid.uuid4()}@example.com"
    test_user_password = "SecurePassword123!"
    
    # 1. Register a test user
    print("\n1. Registering test user...")
    user_data = {
        "email": test_user_email,
        "full_name": "Test User",
        "role": "editor",
        "organization": "Test Organization",
        "password": test_user_password
    }
    
    response = requests.post(f"{base_url}/auth/register", json=user_data)
    assert response.status_code == 200, f"Failed to register test user: {response.text}"
    data = response.json()
    
    # Save token for subsequent tests
    access_token = data["access_token"]
    headers["Authorization"] = f"Bearer {access_token}"
    print("✅ User registration successful")
    
    # 2. Create test files
    print("\n2. Creating test files...")
    
    # Create a DOCX file with formatting
    try:
        from docx import Document
        
        docx_path = "/tmp/test_formatted_document.docx"
        doc = Document()
        
        # Add a title with formatting
        title = doc.add_heading('Formatted Test Document', 0)
        
        # Add a paragraph with bold and italic text
        p = doc.add_paragraph('This document contains ')
        p.add_run('bold text, ').bold = True
        p.add_run('italic text, ').italic = True
        p.add_run('and ')
        p.add_run('bold-italic text.').bold = True
        p.add_run('bold-italic text.').italic = True
        
        # Add a bulleted list
        doc.add_paragraph('List item 1', style='List Bullet')
        doc.add_paragraph('List item 2', style='List Bullet')
        doc.add_paragraph('List item 3', style='List Bullet')
        
        # Add a numbered list
        doc.add_paragraph('Numbered item 1', style='List Number')
        doc.add_paragraph('Numbered item 2', style='List Number')
        doc.add_paragraph('Numbered item 3', style='List Number')
        
        # Add a table
        table = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                table.cell(i, j).text = f"Cell {i+1},{j+1}"
        
        # Save the document
        doc.save(docx_path)
        
        # Create a plain text file for comparison
        txt_path = "/tmp/test_plain_document.txt"
        with open(txt_path, "w") as f:
            f.write("""# Test Plain Document

This is a plain text document with no formatting.

## Section 1
- Item 1
- Item 2
- Item 3

## Section 2
1. First point
2. Second point
3. Third point

## Section 3
This is the final section of our test document.
""")
        
        print("✅ Test files created successfully")
        
        # 3. Test DOCX upload with formatting preservation
        print("\n3. Testing DOCX upload with formatting preservation...")
        
        # Upload the DOCX file
        with open(docx_path, "rb") as f:
            files = {"file": ("test_formatted_document.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                headers={"Authorization": headers["Authorization"]}
            )
        
        assert response.status_code == 200, f"Failed to upload DOCX file: {response.text}"
        data = response.json()
        
        # Verify response structure
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        
        # Verify HTML content generation
        assert "pages" in data["document"], "Missing pages in document"
        assert len(data["document"]["pages"]) > 0, "No pages found in document"
        
        # Check that content is HTML, not plain text
        page_content = data["document"]["pages"][0]["content"]
        assert "<div" in page_content, "Content is not in HTML format"
        
        # Check for formatting elements
        formatting_found = False
        for tag in ["<b>", "<strong>", "<i>", "<em>", "<ul>", "<ol>", "<li>", "<table>"]:
            if tag in page_content.lower():
                formatting_found = True
                break
        
        assert formatting_found, "HTML content does not contain formatting tags"
        
        # Save document ID for retrieval test
        docx_document_id = data["document"]["id"]
        
        # Print some details about the document
        print(f"Document ID: {docx_document_id}")
        print(f"Document title: {data['document']['title']}")
        print(f"Total pages: {data['document']['total_pages']}")
        print(f"Contains formatting: {data['document']['metadata'].get('contains_formatting', False)}")
        
        print("✅ DOCX upload with formatting preservation working")
        
        # 4. Test plain text upload with HTML conversion
        print("\n4. Testing plain text upload with HTML conversion...")
        
        # Upload the text file
        with open(txt_path, "rb") as f:
            files = {"file": ("test_plain_document.txt", f, "text/plain")}
            response = requests.post(
                f"{base_url}/documents/upload",
                files=files,
                headers={"Authorization": headers["Authorization"]}
            )
        
        assert response.status_code == 200, f"Failed to upload text file: {response.text}"
        data = response.json()
        
        # Verify response structure
        assert data["message"] == "Document uploaded and processed successfully", "Unexpected response message"
        assert "document" in data, "Missing document in response"
        
        # Verify HTML content generation for plain text
        assert "pages" in data["document"], "Missing pages in document"
        assert len(data["document"]["pages"]) > 0, "No pages found in document"
        
        # Check that content is HTML, not plain text
        page_content = data["document"]["pages"][0]["content"]
        assert "<div" in page_content, "Content is not in HTML format"
        
        # Save document ID for retrieval test
        txt_document_id = data["document"]["id"]
        
        # Print some details about the document
        print(f"Document ID: {txt_document_id}")
        print(f"Document title: {data['document']['title']}")
        print(f"Total pages: {data['document']['total_pages']}")
        print(f"Contains formatting: {data['document']['metadata'].get('contains_formatting', False)}")
        
        print("✅ Plain text upload with HTML conversion working")
        
        # 5. Test document retrieval with HTML content
        print("\n5. Testing document retrieval with HTML content...")
        
        # Give the server a moment to process the document
        time.sleep(1)
        
        # Retrieve the DOCX document
        response = requests.get(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve DOCX document: {response.text}"
        data = response.json()
        
        # Verify document structure
        assert data["id"] == docx_document_id, "Document ID mismatch"
        assert "pages" in data, "Missing pages in retrieved document"
        assert len(data["pages"]) > 0, "No pages found in retrieved document"
        
        # Check that content is HTML
        page_content = data["pages"][0]["content"]
        assert "<div" in page_content, "Retrieved content is not in HTML format"
        
        # Check for formatting elements
        formatting_found = False
        for tag in ["<b>", "<strong>", "<i>", "<em>", "<ul>", "<ol>", "<li>", "<table>"]:
            if tag in page_content.lower():
                formatting_found = True
                break
        
        assert formatting_found, "Retrieved HTML content does not contain formatting tags"
        
        # Check metadata
        assert "metadata" in data, "Missing metadata in document"
        contains_formatting = data["metadata"].get("contains_formatting", False)
        print(f"Contains formatting flag: {contains_formatting}")
        
        # Print some details about the retrieved document
        print(f"Retrieved document title: {data['title']}")
        print(f"Retrieved document total pages: {data['total_pages']}")
        
        print("✅ Document retrieval with HTML content working")
        
        # 6. Test document page structure
        print("\n6. Testing document page structure...")
        
        # Verify page structure
        assert "pages" in data, "Missing pages in document"
        assert len(data["pages"]) > 0, "No pages found in document"
        
        # Check page structure
        page = data["pages"][0]
        assert "page_number" in page, "Missing page_number in page"
        assert "title" in page, "Missing title in page"
        assert "content" in page, "Missing content in page"
        assert "id" in page, "Missing id in page"
        
        print(f"Page structure: {', '.join(page.keys())}")
        print(f"Page number: {page['page_number']}")
        print(f"Page title: {page['title']}")
        
        print("✅ Document page structure verification working")
        
        # 7. Test updating a document page
        print("\n7. Testing document page update...")
        
        # Update the first page of the DOCX document
        page_update_data = {
            "title": "Updated Page Title",
            "content": "<div style='font-family: Arial; color: blue;'>This is updated content with <b>HTML formatting</b> and <i>styling</i>.</div>"
        }
        
        response = requests.put(
            f"{base_url}/documents/{docx_document_id}/pages/1",
            json=page_update_data,
            headers=headers
        )
        
        assert response.status_code == 200, f"Failed to update document page: {response.text}"
        data = response.json()
        
        # Verify update response
        assert data["message"] == "Page updated successfully", "Unexpected response message"
        
        # Retrieve the document to verify the update
        response = requests.get(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to retrieve updated document: {response.text}"
        data = response.json()
        
        # Check that the page was updated
        assert data["pages"][0]["title"] == "Updated Page Title", "Page title not updated"
        assert "<b>HTML formatting</b>" in data["pages"][0]["content"], "HTML formatting not preserved in update"
        
        print("✅ Document page update with HTML formatting working")
        
        # 8. Test page view tracking
        print("\n8. Testing page view tracking...")
        
        # Track a page view
        view_data = {
            "document_id": docx_document_id,
            "page_number": 1,
            "viewer_info": {
                "ip_address": "192.168.1.1",
                "user_agent": "Test User Agent"
            },
            "time_spent": 60,
            "scroll_depth": 0.8,
            "interactions": ["click", "scroll"]
        }
        
        response = requests.post(f"{base_url}/tracking/page-view", json=view_data)
        assert response.status_code == 200, f"Failed to track page view: {response.text}"
        data = response.json()
        
        # Verify response
        assert data["message"] == "Page view tracked successfully", "Unexpected response message"
        assert "session_id" in data, "Missing session_id in response"
        
        # Save session ID for subsequent tracking
        session_id = data["session_id"]
        
        # Track another page view in the same session
        view_data = {
            "document_id": docx_document_id,
            "page_number": 1,
            "session_id": session_id,
            "viewer_info": {
                "ip_address": "192.168.1.1",
                "user_agent": "Test User Agent"
            },
            "time_spent": 30,
            "scroll_depth": 1.0,
            "interactions": ["click"]
        }
        
        response = requests.post(f"{base_url}/tracking/page-view", json=view_data)
        assert response.status_code == 200, f"Failed to track second page view: {response.text}"
        
        print("✅ Page view tracking working")
        
        # 9. Test page analytics
        print("\n9. Testing page analytics...")
        
        # Get page analytics
        response = requests.get(f"{base_url}/documents/{docx_document_id}/page-analytics", headers=headers)
        assert response.status_code == 200, f"Failed to get page analytics: {response.text}"
        data = response.json()
        
        # Verify analytics structure
        assert "document_id" in data, "Missing document_id in analytics"
        assert "total_pages" in data, "Missing total_pages in analytics"
        assert "page_analytics" in data, "Missing page_analytics in analytics"
        assert "overall_analytics" in data, "Missing overall_analytics in analytics"
        
        # Check page analytics data
        if len(data["page_analytics"]) > 0:
            page_analytics = data["page_analytics"][0]
            print(f"Page analytics for page {page_analytics['page_number']}:")
            print(f"  Total views: {page_analytics['total_views']}")
            print(f"  Average time spent: {page_analytics['average_time_spent']} seconds")
            print(f"  Completion rate: {page_analytics['completion_rate']}%")
        
        print("✅ Page analytics working")
        
        # 10. Clean up
        print("\n10. Cleaning up test documents...")
        
        # Delete the test documents
        response = requests.delete(f"{base_url}/documents/{docx_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to delete DOCX document: {response.text}"
        
        response = requests.delete(f"{base_url}/documents/{txt_document_id}", headers=headers)
        assert response.status_code == 200, f"Failed to delete text document: {response.text}"
        
        # Clean up the test files
        if os.path.exists(docx_path):
            os.remove(docx_path)
        
        if os.path.exists(txt_path):
            os.remove(txt_path)
        
        print("✅ Test cleanup successful")
        
        print("\n✅ All document upload and viewing tests passed successfully!")
    
    except ImportError:
        print("⚠️ python-docx library not available. Installing required dependencies...")
        os.system("pip install python-docx")
        print("Dependencies installed. Please run the test again.")

if __name__ == "__main__":
    test_document_upload_and_viewing()